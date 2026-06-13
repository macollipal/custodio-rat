"""
Custodio RAT Manager — Tema visual común para los .docx
========================================================

Estilo heredado de docs/generar_manual.py y ampliado con:
  - Portada profesional
  - Tabla de contenido (campo TOC)
  - Encabezados y pies de página con número de página
  - Bloques Nota / Advertencia / Riesgo
  - Tablas con estilo "Table Grid" + cabecera sombreada
  - Inserción de figuras Mermaid con caption
  - Numeración correlativa de figuras/tablas por documento
"""

from __future__ import annotations

import os
import subprocess
import tempfile
import uuid
from typing import Iterable

from docx import Document
from docx.document import Document as _Doc
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


# =============================================================================
# CONSTANTES DE MARCA
# =============================================================================

BRAND_NAME = "CUSTODIO"
BRAND_FULL = "Custodio RAT Manager"
BRAND_TAGLINE = "Cumplimiento de la Ley 21.719 de Protección de Datos Personales de Chile"
DOC_VERSION = "v1.0"
DOC_DATE = "Junio 2026"
DOC_AUTHOR = "Equipo de Desarrollo — Custodio"
DOC_COMPANY = "Custodio SpA"
DOC_CLASSIFICATION = "Interno — Confidencial"

# Paleta
COLOR_PRIMARY = RGBColor(0x1F, 0x49, 0x7D)   # azul institucional
COLOR_SECONDARY = RGBColor(0x2E, 0x74, 0xB5) # azul claro
COLOR_TEXT = RGBColor(0x40, 0x40, 0x40)     # gris oscuro para H3
COLOR_BODY = RGBColor(0x00, 0x00, 0x00)
COLOR_MUTED = RGBColor(0x80, 0x80, 0x80)
COLOR_HEADER_BG = "D9E2F3"   # azul muy claro (header de tabla)
COLOR_NOTA_BG = "F2F7FC"     # azul muy suave
COLOR_ADV_BG = "FFF8E5"      # amarillo suave
COLOR_RIESGO_BG = "FBEAEA"   # rojo suave
COLOR_BORDER_NOTA = "2E74B5"
COLOR_BORDER_ADV = "BF8F00"
COLOR_BORDER_RIESGO = "C00000"


# =============================================================================
# CREACIÓN DEL DOCUMENTO BASE
# =============================================================================

def new_document() -> _Doc:
    """Crea un documento con los estilos por defecto del tema Custodio."""
    doc = Document()

    # --- Márgenes (Carta) ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_height = Cm(27.94)
        section.page_width = Cm(21.59)

    # --- Estilo Normal ---
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_BODY
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # --- Heading 1 ---
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_PRIMARY
    pf1 = h1.paragraph_format
    pf1.space_before = Pt(18)
    pf1.space_after = Pt(6)
    pf1.keep_with_next = True

    # --- Heading 2 ---
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_SECONDARY
    pf2 = h2.paragraph_format
    pf2.space_before = Pt(12)
    pf2.space_after = Pt(4)
    pf2.keep_with_next = True

    # --- Heading 3 ---
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_TEXT
    pf3 = h3.paragraph_format
    pf3.space_before = Pt(8)
    pf3.space_after = Pt(2)
    pf3.keep_with_next = True

    # --- Heading 4 ---
    h4 = doc.styles["Heading 4"]
    h4.font.name = "Calibri"
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = COLOR_TEXT

    return doc


# =============================================================================
# ENCABEZADO Y PIE DE PÁGINA
# =============================================================================

def _add_field(paragraph, field_code: str) -> None:
    """Inserta un campo (PAGE, NUMPAGES, TOC, etc.) en el párrafo."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def add_header_footer(doc: _Doc, doc_title: str) -> None:
    """Aplica encabezado y pie de página con número a todas las secciones."""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # --- Encabezado (no aparece en primera página = portada) ---
        header = section.header
        p_hdr = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p_hdr.text = ""
        # Tabla de 2 columnas para layout limpio
        tbl = header.add_table(rows=1, cols=2, width=Cm(17.59))
        tbl.autofit = False
        tbl.columns[0].width = Cm(10.0)
        tbl.columns[1].width = Cm(7.59)
        # Celda izquierda: isotipo + nombre doc
        c0 = tbl.rows[0].cells[0]
        c0.width = Cm(10.0)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0a = p0.add_run("▣  CUSTODIO  ")
        r0a.bold = True
        r0a.font.size = Pt(9)
        r0a.font.color.rgb = COLOR_PRIMARY
        r0b = p0.add_run("·  " + doc_title)
        r0b.font.size = Pt(9)
        r0b.font.color.rgb = COLOR_TEXT
        # Celda derecha: versión
        c1 = tbl.rows[0].cells[1]
        c1.width = Cm(7.59)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(DOC_VERSION + "  ·  " + DOC_DATE)
        r1.font.size = Pt(9)
        r1.font.color.rgb = COLOR_MUTED
        # Línea inferior del encabezado
        _add_bottom_border(p0, color="1F497D", size_pt=0.75)

        # --- Pie de página ---
        footer = section.footer
        p_ftr = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p_ftr.text = ""
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Tabla 1x3: izq / centro / der
        ftbl = footer.add_table(rows=1, cols=3, width=Cm(17.59))
        ftbl.autofit = False
        ftbl.columns[0].width = Cm(6.0)
        ftbl.columns[1].width = Cm(5.59)
        ftbl.columns[2].width = Cm(6.0)
        # Izq
        f0 = ftbl.rows[0].cells[0]
        f0.width = Cm(6.0)
        fp0 = f0.paragraphs[0]
        fp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rfp0 = fp0.add_run(f"{BRAND_FULL} · Ley 21.719")
        rfp0.font.size = Pt(8)
        rfp0.font.color.rgb = COLOR_MUTED
        # Centro: "Página X de Y"
        f1 = ftbl.rows[0].cells[1]
        f1.width = Cm(5.59)
        fp1 = f1.paragraphs[0]
        fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rfp1a = fp1.add_run("Página ")
        rfp1a.font.size = Pt(8)
        rfp1a.font.color.rgb = COLOR_MUTED
        _add_field(fp1, "PAGE")
        rfp1b = fp1.add_run(" de ")
        rfp1b.font.size = Pt(8)
        rfp1b.font.color.rgb = COLOR_MUTED
        _add_field(fp1, "NUMPAGES")
        for r in fp1.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_MUTED
        # Der: clasificación
        f2 = ftbl.rows[0].cells[2]
        f2.width = Cm(6.0)
        fp2 = f2.paragraphs[0]
        fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rfp2 = fp2.add_run(f"{DOC_CLASSIFICATION}  ·  " + doc_title)
        rfp2.font.size = Pt(8)
        rfp2.font.color.rgb = COLOR_MUTED


def _add_bottom_border(paragraph, color="1F497D", size_pt=0.75) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_pt * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    # Eliminar border previo si existiera
    for child in pBdr.findall(qn("w:bottom")):
        pBdr.remove(child)
    pBdr.append(bottom)


# =============================================================================
# PORTADA
# =============================================================================

def add_cover(doc: _Doc, title: str, subtitle: str, code: str) -> None:
    """Inserta una portada profesional al inicio del documento."""
    # Espaciado vertical inicial
    for _ in range(3):
        doc.add_paragraph()

    # Banda azul superior
    band_top = doc.add_paragraph()
    band_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_border_box(band_top, color="1F497D", size_pt=2.0,
                              background="1F497D")
    band_top_run = band_top.add_run("  ▣  CUSTODIO  ·  " + BRAND_FULL.upper() + "  ")
    band_top_run.font.size = Pt(12)
    band_top_run.font.bold = True
    band_top_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Código del documento
    p_code = doc.add_paragraph()
    p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_code = p_code.add_run(code)
    r_code.font.size = Pt(11)
    r_code.font.color.rgb = COLOR_MUTED
    r_code.italic = True
    p_code.paragraph_format.space_before = Pt(36)
    p_code.paragraph_format.space_after = Pt(6)

    # Título principal
    for _ in range(2):
        doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title)
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY
    p_title.paragraph_format.space_after = Pt(12)

    # Subtítulo
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(subtitle)
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = COLOR_SECONDARY
    r_sub.italic = True
    p_sub.paragraph_format.space_after = Pt(24)

    # Línea separadora
    for _ in range(2):
        doc.add_paragraph()
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bottom_border(p_sep, color="2E74B5", size_pt=1.5)
    r_sep = p_sep.add_run("  ")
    r_sep.font.size = Pt(2)

    # Metadatos
    for _ in range(2):
        doc.add_paragraph()
    meta_lines = [
        ("Versión", DOC_VERSION),
        ("Fecha de emisión", DOC_DATE),
        ("Autor", DOC_AUTHOR),
        ("Empresa responsable", DOC_COMPANY),
        ("Clasificación", DOC_CLASSIFICATION),
        ("Producto", BRAND_FULL),
        ("Marco regulatorio", "Ley N.° 21.719 — Chile"),
    ]
    for label, value in meta_lines:
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p_m.add_run(label + ":  ")
        rl.font.size = Pt(11)
        rl.font.bold = True
        rl.font.color.rgb = COLOR_PRIMARY
        rv = p_m.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = COLOR_TEXT
        p_m.paragraph_format.space_after = Pt(2)

    # Pie de portada
    for _ in range(2):
        doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_border_box(p_foot, color="1F497D", size_pt=2.0,
                              background="1F497D")
    rf = p_foot.add_run("  " + BRAND_TAGLINE.upper() + "  ")
    rf.font.size = Pt(10)
    rf.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rf.bold = True

    # Salto de página
    doc.add_page_break()


def _add_paragraph_border_box(paragraph, color="1F497D", size_pt=1.0,
                              background=None) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(int(size_pt * 8)))
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    if background:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), background)
        pPr.append(shd)


# =============================================================================
# PÁGINAS DE CONTROL DE VERSIONES Y TOC
# =============================================================================

def add_version_control(doc: _Doc, code: str, doc_title: str,
                        changes: list[tuple[str, str, str]] | None = None) -> None:
    """Añade la página de control de versiones."""
    h = doc.add_heading("Control de versiones", level=1)

    if changes is None:
        changes = [
            ("1.0", DOC_DATE, "Creación inicial del documento."),
        ]

    tbl = doc.add_table(rows=1 + len(changes), cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, name in enumerate(["Versión", "Fecha", "Autor", "Cambios"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(name)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr[i], COLOR_PRIMARY)
        # Anchos: Versión 1.5, Fecha 2.5, Autor 3.0, Cambios 10.59
        col_widths = [Cm(1.5), Cm(2.5), Cm(3.0), Cm(10.59)]
        hdr[i].width = col_widths[i]

    for ri, (v, d, c) in enumerate(changes, start=1):
        row = tbl.rows[ri].cells
        row[0].text = v
        row[1].text = d
        row[2].text = DOC_AUTHOR
        row[3].text = ""
        p = row[3].paragraphs[0]
        r = p.add_run(c)
        r.font.size = Pt(10)
        # Subrayar la última entrada (la más reciente)
        if ri == len(changes):
            r.underline = True
        for ci in range(3):
            row[ci].width = [Cm(1.5), Cm(2.5), Cm(3.0)][ci]
            for p in row[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # Estado del documento
    doc.add_paragraph()
    h2 = doc.add_heading("Estado del documento", level=2)
    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Código", code),
        ("Nombre", doc_title),
        ("Producto", BRAND_FULL),
        ("Marco regulatorio", "Ley N.° 21.719 de Protección de Datos Personales (Chile)"),
    ]
    tbl2.rows[0].cells[0].text = ""
    p = tbl2.rows[0].cells[0].paragraphs[0]
    r = p.add_run("Campo")
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_shading(tbl2.rows[0].cells[0], COLOR_PRIMARY)
    p2 = tbl2.rows[0].cells[1].paragraphs[0]
    r2 = p2.add_run("Valor")
    r2.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_shading(tbl2.rows[0].cells[1], COLOR_PRIMARY)
    for i, (k, v) in enumerate(rows_data[1:], start=1):
        tbl2.rows[i].cells[0].text = k
        tbl2.rows[i].cells[1].text = v
        for c in tbl2.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_page_break()


def add_toc(doc: _Doc) -> None:
    """Inserta una tabla de contenido (campo TOC). Word la actualiza al abrir."""
    doc.add_heading("Tabla de contenido", level=1)
    p = doc.add_paragraph()
    _add_field(p, r'TOC \o "1-3" \h \z \u')
    # Nota al usuario
    note = doc.add_paragraph()
    rn = note.add_run(
        "Sugerencia: en Microsoft Word, presione F9 sobre la tabla para "
        "actualizar el contenido y los números de página.")
    rn.italic = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = COLOR_MUTED
    doc.add_page_break()


def _set_cell_shading(cell, color) -> None:
    """Aplica un color de fondo a una celda. Acepta hex string o RGBColor."""
    if hasattr(color, "_color"):
        # RGBColor: convertir a hex
        c = color._color
        if hasattr(c, "rgb"):
            hex_str = str(c.rgb)
        else:
            hex_str = str(c)
    else:
        hex_str = str(color)
    hex_str = hex_str.upper().lstrip("#")
    # rgb='1F497D' -> ya viene en hex sin #. Asegurar 6 chars.
    if len(hex_str) > 6:
        hex_str = hex_str[-6:]
    tcPr = cell._tc.get_or_add_tcPr()
    # Eliminar shading previo
    for el in tcPr.findall(qn("w:shd")):
        tcPr.remove(el)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tcPr.append(shd)


# =====================================================================
# UTILIDADES PARA TABLAS
# =====================================================================

def add_styled_table(doc: _Doc, headers: list[str], rows: list[list[str]],
                     col_widths_cm: list[float] | None = None,
                     header_bg: str = "D9E2F3",
                     first_col_bold: bool = False,
                     underline_new: bool = False) -> None:
    """Añade una tabla con estilo Table Grid y cabecera sombreada.
    Si underline_new=True, el texto envuelto en guiones bajos _texto_ se subrayará."""
    if not rows and not headers:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    # Cabecera
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_PRIMARY
        _set_cell_shading(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Datos
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            text = str(val) if val is not None else ""
            # Si underline_new y el texto empieza y termina con _, aplicar underline
            if underline_new and text.startswith("_") and text.endswith("_"):
                text = text[1:-1]  # Quitar los _
                r = p.add_run(text)
                r.font.size = Pt(10)
                r.underline = True
            else:
                r = p.add_run(text)
                r.font.size = Pt(10)
            if first_col_bold and ci == 0:
                r.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_caption_table(doc: _Doc, caption: str,
                      counter: list[int], label: str = "Tabla") -> None:
    """Añade un caption tipo 'Tabla N — descripción'."""
    counter[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label} {counter[0]} — ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.italic = True
    r1.font.color.rgb = COLOR_PRIMARY
    r2 = p.add_run(caption)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = COLOR_TEXT


# =============================================================================
# BLOQUES NOTA / ADVERTENCIA / RIESGO
# =============================================================================

def add_note(doc: _Doc, title: str, body: str,
             border: str = COLOR_BORDER_NOTA,
             background: str = COLOR_NOTA_BG,
             icon: str = "ℹ") -> None:
    p = doc.add_paragraph()
    _add_paragraph_border_box(p, color=border, size_pt=2.0, background=background)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    r1 = p.add_run(f"{icon}  {title}.  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = COLOR_PRIMARY if border == COLOR_BORDER_NOTA else (
        RGBColor(0xBF, 0x8F, 0x00) if border == COLOR_BORDER_ADV else
        RGBColor(0xC0, 0x00, 0x00))
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = COLOR_TEXT


def add_warning(doc: _Doc, title: str, body: str) -> None:
    add_note(doc, title, body,
             border=COLOR_BORDER_ADV, background=COLOR_ADV_BG, icon="⚠")


def add_risk(doc: _Doc, title: str, body: str) -> None:
    add_note(doc, title, body,
             border=COLOR_BORDER_RIESGO, background=COLOR_RIESGO_BG, icon="⛔")


# =============================================================================
# RENDERIZADO DE MERMAID → PNG
# =============================================================================

MERMAID_CACHE: dict[str, str] = {}


def render_mermaid(mermaid_code: str, output_dir: str, name_hint: str = "fig",
                   width_px: int = 1600, height_px: int = 900) -> str:
    """Renderiza un bloque Mermaid a PNG (2x) usando mermaid-cli (npx).
    Devuelve la ruta absoluta del PNG. Usa caché por hash del código.
    """
    cache_key = f"{name_hint}_{hash(mermaid_code)}"
    if cache_key in MERMAID_CACHE:
        return MERMAID_CACHE[cache_key]

    os.makedirs(output_dir, exist_ok=True)
    mmd_path = os.path.join(tempfile.gettempdir(),
                            f"mmd_{uuid.uuid4().hex[:8]}.mmd")
    png_path = os.path.join(output_dir, f"{cache_key}.png")

    with open(mmd_path, "w", encoding="utf-8") as f:
        f.write(mermaid_code)

    # Mermaid CLI (mmdc) — npx -y @mermaid-js/mermaid-cli
    cmd = [
        "npx.cmd", "-y", "@mermaid-js/mermaid-cli@10.9.1",
        "-i", mmd_path,
        "-o", png_path,
        "-w", str(width_px),
        "-H", str(height_px),
        "-b", "white",
        "-s", "2",
    ]
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=180,
            shell=True,
        )
        if result.returncode != 0 or not os.path.exists(png_path):
            err = (result.stderr or result.stdout or "")[:2000]
            raise RuntimeError(f"mermaid-cli falló: {err}")
    except Exception as e:
        # Si falla, devolvemos un PNG placeholder (no cortar la generación)
        placeholder = os.path.join(output_dir, f"{cache_key}_ERROR.png")
        _write_placeholder_png(placeholder, f"Error Mermaid: {e}")
        png_path = placeholder

    try:
        os.remove(mmd_path)
    except OSError:
        pass

    MERMAID_CACHE[cache_key] = png_path
    return png_path


def _write_placeholder_png(path: str, text: str) -> None:
    """Escribe un PNG mínimo (1x1 blanco) si Mermaid falla — evita cortar."""
    # PNG 1x1 blanco, base64
    import base64
    b64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNg"
        "AAIAAAUAAeImBZsAAAAASUVORK5CYII="
    )
    with open(path, "wb") as f:
        f.write(base64.b64decode(b64))
    # Y un txt al lado con el mensaje (UTF-8, sin mojibake)
    try:
        with open(path + ".txt", "w", encoding="utf-8") as f:
            f.write(text)
    except OSError:
        pass


def add_figure(doc: _Doc, mermaid_code: str, caption: str,
               output_dir: str, counter: list[int],
               width_inches: float = 6.5,
               name_hint: str = "fig",
               max_height_inches: float = 5.0) -> None:
    """Renderiza un Mermaid a PNG y lo inserta con caption."""
    png_path = render_mermaid(mermaid_code, output_dir, name_hint=name_hint)
    try:
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        # Intentar escalar por aspect ratio si se conoce
        from PIL import Image
        try:
            with Image.open(png_path) as im:
                w, h = im.size
            if w and h:
                ratio = h / w
                target_h = width_inches * ratio
                if target_h > max_height_inches:
                    width_inches = max_height_inches / ratio
                run.add_picture(png_path, width=Inches(width_inches))
            else:
                run.add_picture(png_path, width=Inches(width_inches))
        except Exception:
            run.add_picture(png_path, width=Inches(width_inches))
    except Exception as e:
        # Si python-docx no logra insertar, registrar como nota
        add_warning(doc, "Figura omitida",
                    f"No se pudo insertar la imagen: {e}. Revisar: {name_hint}")
    counter[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run(f"Figura {counter[0]} — ")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = COLOR_PRIMARY
    r2 = p.add_run(caption)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = COLOR_TEXT


# =============================================================================
# HELPER PARA LISTAS
# =============================================================================

def add_bullet(doc: _Doc, text: str, level: int = 0,
               bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)


def add_numbered(doc: _Doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)


def add_paragraph(doc: _Doc, text: str, italic: bool = False,
                  bold: bool = False, alignment=None) -> None:
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    r = p.add_run(text)
    r.font.size = Pt(11)
    if italic:
        r.italic = True
    if bold:
        r.bold = True


def add_kv_table(doc: _Doc, items: list[tuple[str, str]]) -> None:
    """Tabla simple de 2 columnas Clave / Valor."""
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(items):
        tbl.rows[i].cells[0].text = ""
        p = tbl.rows[i].cells[0].paragraphs[0]
        r = p.add_run(k)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_PRIMARY
        _set_cell_shading(tbl.rows[i].cells[0], "F2F2F2")
        tbl.rows[i].cells[1].text = ""
        p2 = tbl.rows[i].cells[1].paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.size = Pt(10)


# =============================================================================
# CIERRES COMUNES: APÉNDICES FINALES
# =============================================================================

def add_open_questions(doc: _Doc, questions: list[str]) -> None:
    doc.add_heading("Apéndice A — Vacíos de información", level=1)
    add_paragraph(doc,
                  "Las siguientes preguntas quedan abiertas al equipo de producto "
                  "o al Product Owner para ser resueltas en próximas iteraciones:",
                  italic=True)
    for q in questions:
        add_bullet(doc, q)


def add_risks_appendix(doc: _Doc, risks: list[tuple[str, str, str]]) -> None:
    """risks: lista de (ID, descripción, severidad)."""
    doc.add_heading("Apéndice B — Riesgos identificados", level=1)
    add_paragraph(doc,
                  "Riesgos detectados durante la auditoría documental y técnica:",
                  italic=True)
    rows = [[r[0], r[1], r[2]] for r in risks]
    add_caption_table(doc, "Riesgos identificados durante la auditoría", [0], "Tabla")
    add_styled_table(doc, ["ID", "Descripción", "Severidad"], rows,
                     col_widths_cm=[2.0, 13.0, 2.5])


def add_id_glossary(doc: _Doc, glossary: list[tuple[str, str, str]]) -> None:
    """glossary: (ID, nombre, descripción)."""
    doc.add_heading("Apéndice C — Índice de identificadores", level=1)
    add_paragraph(doc,
                  "Identificadores únicos referenciados a lo largo del documento:",
                  italic=True)
    rows = [[g[0], g[1], g[2]] for g in glossary]
    add_styled_table(doc, ["ID", "Nombre", "Descripción"], rows,
                     col_widths_cm=[2.5, 5.0, 10.0], first_col_bold=True)


def add_final_note(doc: _Doc) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—  Fin del documento  —")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_MUTED
