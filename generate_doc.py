"""
Genera el documento Word con el Plan de Mejoras de Seguridad — Custodio RAT
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Title
title = doc.add_heading('Plan de Mejoras de Seguridad — Custodio RAT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run('MFA/2FA · Backups Automatizados · Cifrado de Datos en Reposo')
sub_run.italic = True
sub_run.font.size = Pt(12)

doc.add_paragraph()

# ============= RESUMEN EJECUTIVO =============
doc.add_heading('Resumen Ejecutivo', 1)
doc.add_paragraph(
    'Los 3 puntos propuestos (MFA/2FA, backups automatizados, cifrado de datos en reposo) '
    'son componentes críticos de seguridad empresarial. La implementación es viable en 3 fases, '
    'sin romper compatibilidad con la arquitectura actual.'
)
doc.add_paragraph(
    'Custodio RAT implementa controles de seguridad fundamentales correctamente: hashing bcrypt, '
    'JWT con expiración, rate limiting, httpOnly cookies, y ORM contra inyección SQL. '
    'Sin embargo, existen debilidades críticas que deben abordarse antes de producción, '
    'especialmente en la gestión de secretos, el almacenamiento de tokens en localStorage, '
    'y la ausencia de MFA.'
)
p = doc.add_paragraph()
p.add_run('Clasificación actual: ').bold = True
p.add_run('Producción Inicial — Cumple requisitos mínimos pero no está listo para manejo '
         'de datos sensibles regulados sin correcciones.')

# ============= PUNTO 1: MFA =============
doc.add_heading('Punto 1: Autenticación Multifactor (MFA / 2FA)', 1)

doc.add_heading('1.1 Análisis del problema', 2)
doc.add_paragraph(
    'La autenticación actual es single-factor (usuario + contraseña). '
    'Para una plataforma que maneja datos regulados por la Ley 21.719, esto es insuficiente.'
)

doc.add_heading('1.2 Opciones de implementación', 2)

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Opción'
hdr[1].text = 'Descripción'
hdr[2].text = 'Pros'
hdr[3].text = 'Contras'
hdr[4].text = 'Tiempo'

opciones = [
    ('TOTP (RFC 6238)', 'Google Authenticator, Authy, 1Password',
     'Sin costos, estándar maduro, sin dependencias externas',
     'Usuario necesita app, soporte al usuario', '2-3 semanas'),
    ('WebAuthn / FIDO2', 'Llaves físicas (YubiKey) o passkeys',
     'Máxima seguridad, no susceptible a phishing',
     'Costo de llaves, requiere HTTPS robusto', '4-6 semanas'),
    ('SMS OTP', 'Código por SMS',
     'Familiar para usuarios',
     'Costo SMS, susceptible a SIM swap', '1-2 semanas'),
    ('Email OTP', 'Código por email',
     'Bajo costo, no requiere app',
     'Depende de email, susceptible a robo', '1 semana'),
]
for opt in opciones:
    row = table.add_row().cells
    for i, val in enumerate(opt):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('1.3 Recomendación', 2)
p = doc.add_paragraph()
p.add_run('TOTP como base + Email OTP como fallback.').bold = True
p.add_run(' Es el estándar de la industria (NIST 800-63B), '
         'no depende de proveedores externos, y es compatible con apps gratuitas.')

doc.add_heading('1.4 Diseño arquitectónico', 2)
doc.add_paragraph(
    'Frontend (Next.js):\n'
    '  - Pantalla de configuración TOTP\n'
    '  - Pantalla de verificación 6 dígitos\n'
    '  - QR code con secret\n\n'
    'Backend FastAPI:\n'
    '  - POST /auth/mfa/setup → genera secret\n'
    '  - POST /auth/mfa/verify → activa MFA\n'
    '  - POST /auth/mfa/disable\n'
    '  - POST /auth/login acepta MFA token\n'
    '  - GET  /auth/mfa/backup-codes\n\n'
    'Database:\n'
    '  - users.mfa_secret (encrypted)\n'
    '  - users.mfa_enabled (boolean)\n'
    '  - users.mfa_backup_codes (hashed)\n'
    '  - audit_log entries'
)

doc.add_heading('1.5 Cambios necesarios', 2)
doc.add_paragraph('Backend:', style='List Bullet').text = 'app/models/user.py — agregar campos mfa_secret, mfa_enabled, mfa_backup_codes'
doc.add_paragraph('app/schemas/user.py — schemas para MfaSetupRequest, MfaVerifyRequest')
doc.add_paragraph('app/core/mfa.py (nuevo) — generación de secret, QR code, verificación TOTP')
doc.add_paragraph('app/routes/auth.py — endpoints MFA + integración con login')
doc.add_paragraph('app/core/security.py — cifrar mfa_secret antes de guardar')

doc.add_paragraph('Frontend:', style='List Bullet').text = 'app/(app)/configuracion/page.tsx o nueva app/(app)/perfil/mfa/page.tsx'
doc.add_paragraph('app/login/page.tsx — pantalla de verificación MFA')
doc.add_paragraph('context/AppContext.tsx — manejar estado mfa_required')

doc.add_paragraph('Dependencias:', style='List Bullet').text = 'pyotp — generación y verificación TOTP'
doc.add_paragraph('qrcode[pil] — generación de QR')
doc.add_paragraph('cryptography — cifrado de secretos')

doc.add_heading('1.6 Compliance', 2)
doc.add_paragraph('Cumple NIST 800-63B AAL2 (Authenticator Assurance Level 2)', style='List Bullet')
doc.add_paragraph('Recomendado por Ley 21.719 Art. 16 BIS para datos sensibles')
doc.add_paragraph('Compatible con OWASP ASVS V2.5.1')

# ============= PUNTO 2: BACKUPS =============
doc.add_heading('Punto 2: Plan de Backup de Base de Datos', 1)

doc.add_heading('2.1 Análisis del problema', 2)
doc.add_paragraph(
    'La base de datos es PostgreSQL en Neon. Neon ofrece backups automáticos PITR '
    '(Point-in-Time Recovery) de 7 días en planes gratuitos y 30 días en planes Pro. '
    'Pero no hay evidencia de backups automatizados propios ni verificación de integridad.'
)

doc.add_heading('2.2 Estrategia de backup recomendada (3-2-1)', 2)
p = doc.add_paragraph()
p.add_run('3 copias de la BD · 2 medios diferentes · 1 offsite').italic = True

table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Light Grid Accent 1'
hdr = table2.rows[0].cells
hdr[0].text = 'Tipo'
hdr[1].text = 'Frecuencia'
hdr[2].text = 'Retención'
hdr[3].text = 'Storage'
hdr[4].text = 'Costo estimado'

backups = [
    ('Snapshot completo', 'Diario (03:00 AM)', '30 días', 'S3 + cifrado AES-256', '~$10-20/mes'),
    ('PITR (Neon nativo)', 'Continuo', '7 días', 'Neon Pro', 'Incluido en Neon'),
    ('Backup mensual', 'Mensual', '12 meses', 'Glacier S3', '~$2/mes'),
    ('Backup offsite', 'Semanal', '90 días', 'Google Cloud Storage', '~$5/mes'),
]
for b in backups:
    row = table2.add_row().cells
    for i, val in enumerate(b):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('2.3 Diseño arquitectónico', 2)
doc.add_paragraph(
    'Neon PostgreSQL (producción) — BD principal\n'
    '        ↓ pg_dump (cron 03:00)\n'
    'Vercel Cron / GitHub Actions — Job automatizado\n'
    '        ↓ gzip + cifrado AES-256\n'
    'AWS S3 / GCS — Almacenamiento (3-2-1 backup)\n'
    '        ↓ Mensual → Glacier\n'
    'Cold Storage — Retención larga'
)

doc.add_heading('2.4 Implementación', 2)
p = doc.add_paragraph()
p.add_run('Opción A: GitHub Actions (recomendado)').bold = True
doc.add_paragraph(
    'name: DB Backup\n'
    'on:\n'
    '  schedule: { cron: \'0 3 * * *\' }  # 3 AM diario\n'
    '  workflow_dispatch:\n'
    'jobs:\n'
    '  backup:\n'
    '    runs-on: ubuntu-latest\n'
    '    steps:\n'
    '      - uses: actions/checkout@v4\n'
    '      - name: Install pg_dump\n'
    '        run: sudo apt-get install postgresql-client\n'
    '      - name: Backup\n'
    '        env:\n'
    '          DATABASE_URL: ${{ secrets.NEON_DATABASE_URL }}\n'
    '          BACKUP_ENCRYPTION_KEY: ${{ secrets.BACKUP_KEY }}\n'
    '        run: |\n'
    '          pg_dump "$DATABASE_URL" | gzip | \\\n'
    '          openssl enc -aes-256-gcm -salt -pass "$BACKUP_ENCRYPTION_KEY" \\\n'
    '          > backup-$(date +%Y%m%d).sql.gz.enc\n'
    '      - name: Upload to S3\n'
    '        uses: aws-actions/configure-aws-credentials@v4\n'
    '        with:\n'
    '          aws-access-key-id: ${{ secrets.AWS_ACCESS_KEY_ID }}\n'
    '          aws-secret-access-key: ${{ secrets.AWS_SECRET_ACCESS_KEY }}\n'
    '          aws-region: us-east-1\n'
    '      - run: aws s3 cp backup-*.sql.gz.enc s3://custodio-backups/daily/\n'
    '      - name: Verify integrity\n'
    '        run: |\n'
    '          openssl enc -d -aes-256-gcm -in backup-*.sql.gz.enc -pass "$BACKUP_ENCRYPTION_KEY" | \\\n'
    '          gunzip | head -c 100 | grep -q "PostgreSQL"'
)

doc.add_heading('2.5 Verificación de integridad', 2)
doc.add_paragraph('Restore test mensual — descargar backup, descifrar, restaurar en BD staging, verificar', style='List Bullet')
doc.add_paragraph('Checksum SHA-256 de cada backup al guardar')
doc.add_paragraph('Alerta si backup diario falla')

doc.add_heading('2.6 Compliance', 2)
doc.add_paragraph('Cumple ISO 27001 A.12.3.1 (backup de información)', style='List Bullet')
doc.add_paragraph('Cumple Ley 21.719 para retención de evidencia regulatoria')
p = doc.add_paragraph()
p.add_run('RTO objetivo: ').bold = True
p.add_run('4 horas (restore completo)')
p = doc.add_paragraph()
p.add_run('RPO objetivo: ').bold = True
p.add_run('24 horas (backup diario)')

# ============= PUNTO 3: CIFRADO =============
doc.add_heading('Punto 3: Cifrado de Datos en Reposo', 1)

doc.add_heading('3.1 Análisis del problema', 2)
p = doc.add_paragraph()
p.add_run('Actualmente:\n').bold = True
p.add_run('✅ Conexión DB cifrada con SSL (tránsito)\n')
p.add_run('✅ Passwords hasheadas con bcrypt\n')
p.add_run('❌ Datos personales (RUT, email, nombre) almacenados en texto plano\n')
p.add_run('❌ Campos sensibles (datos de salud, financieros) en texto plano')

doc.add_heading('3.2 Estrategia de cifrado recomendada (híbrido)', 2)

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
hdr = table3.rows[0].cells
hdr[0].text = 'Capa'
hdr[1].text = 'Tecnología'
hdr[2].text = 'Datos protegidos'

capas = [
    ('Nivel aplicación (columna)', 'cryptography (Fernet)', 'RUT, email, nombre titular, datos sensibles'),
    ('Nivel DB (TDE)', 'Neon nativo + AWS KMS', 'BD completa en reposo'),
    ('Nivel backup', 'AES-256-GCM', 'Archivos de backup'),
    ('Nivel secreto', 'Sealed Secrets / Vault', 'Secret keys, credenciales SMTP'),
]
for c in capas:
    row = table3.add_row().cells
    for i, val in enumerate(c):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('3.3 Diseño de cifrado a nivel columna', 2)
doc.add_paragraph(
    'Dato en frontend → POST /api → Backend\n'
    '                              ↓\n'
    '                    Cifrado con Fernet (AES-128-CBC + HMAC)\n'
    '                              ↓\n'
    '                    Almacenado en DB como BYTEA/VARCHAR'
)
p = doc.add_paragraph()
p.add_run('Implementación:').bold = True
doc.add_paragraph(
    'from cryptography.fernet import Fernet\n'
    'from app.core.config import settings\n\n'
    'class FieldEncryption:\n'
    '    def __init__(self):\n'
    '        key = settings.FERNET_KEY.encode()  # 32 url-safe base64 bytes\n'
    '        self.fernet = Fernet(key)\n'
    '    \n'
    '    def encrypt(self, plaintext: str) -> str:\n'
    '        if plaintext is None:\n'
    '            return None\n'
    '        return self.fernet.encrypt(plaintext.encode()).decode()\n'
    '    \n'
    '    def decrypt(self, ciphertext: str) -> str:\n'
    '        if ciphertext is None:\n'
    '            return None\n'
    '        return self.fernet.decrypt(ciphertext.encode()).decode()'
)

doc.add_heading('3.4 Campos a cifrar', 2)
p = doc.add_paragraph()
p.add_run('Tabla users:').bold = True
doc.add_paragraph('email 🔐', style='List Bullet')
doc.add_paragraph('full_name 🔐')
doc.add_paragraph('hashed_password (ya está hasheado con bcrypt) ✅')

p = doc.add_paragraph()
p.add_run('Tabla tkt_solicitud_derecho:').bold = True
doc.add_paragraph('titular_nombre 🔐', style='List Bullet')
doc.add_paragraph('titular_email 🔐')
doc.add_paragraph('titular_rut 🔐')
doc.add_paragraph('descripcion 🔐')

p = doc.add_paragraph()
p.add_run('Tabla rats:').bold = True
doc.add_paragraph('Campos que contengan datos sensibles o personales identificables', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Tabla tkt_notas:').bold = True
doc.add_paragraph('nota 🔐 (puede contener info sensible)', style='List Bullet')

doc.add_heading('3.5 Gestión de claves (KMS)', 2)
doc.add_paragraph(
    'Jerarquía de claves:\n'
    'AWS KMS Master Key (rotación automática anual)\n'
    '  ↓\n'
    'Fernet Data Encryption Key (DEK) - por campo/tabla\n'
    '  ↓\n'
    'Cifrado de columnas'
)
p = doc.add_paragraph()
p.add_run('Opciones de KMS:').bold = True
doc.add_paragraph('AWS KMS (~$1/key/mes) — recomendado para producción', style='List Bullet')
doc.add_paragraph('HashiCorp Vault — open source, más complejo')
doc.add_paragraph('Sealed Secrets de Bitnami — para Kubernetes (no aplica aún)')

doc.add_heading('3.6 Rotación de claves', 2)
doc.add_paragraph('Master key: rotación anual automática', style='List Bullet')
doc.add_paragraph('Data encryption keys: rotación trimestral')
doc.add_paragraph('Re-cifrado: cuando se rota una DEK, hay que re-cifrar todas las columnas (proceso batch)')

doc.add_heading('3.7 Compliance', 2)
doc.add_paragraph('Cumple GDPR Art. 32 (seguridad del tratamiento)', style='List Bullet')
doc.add_paragraph('Cumple Ley 21.719 Art. 14 (medidas de seguridad)')
doc.add_paragraph('Cumple PCI-DSS 3.4 (protección de datos en reposo)')
doc.add_paragraph('Cumple ISO 27001 A.10.1.1')

# ============= ROADMAP =============
doc.add_heading('Roadmap de Implementación', 1)

doc.add_heading('Fase 1 — Quick wins (1-2 semanas)', 2)
doc.add_paragraph('TOTP MFA con pyotp y qrcode', style='List Bullet')
doc.add_paragraph('Backup diario automatizado a S3 con cifrado AES-256')
doc.add_paragraph('Cifrado Fernet de campos sensibles más críticos (users.email, tkt_solicitud_derecho.titular_rut)')

doc.add_heading('Fase 2 — Refuerzo (3-4 semanas)', 2)
doc.add_paragraph('Cifrado del resto de campos personales', style='List Bullet')
doc.add_paragraph('Backup verificado con restore test mensual')
doc.add_paragraph('MFA obligatorio para rol superadmin y admin_empresa')
doc.add_paragraph('Rotación de claves implementada')

doc.add_heading('Fase 3 — Empresa (6-8 semanas)', 2)
doc.add_paragraph('WebAuthn/FIDO2 como segunda opción MFA', style='List Bullet')
doc.add_paragraph('Backup offsite en segunda región')
doc.add_paragraph('AWS KMS para gestión centralizada de claves')
doc.add_paragraph('Auditoría de cifrado anual')

# ============= COSTOS =============
doc.add_heading('Estimación de Costos', 1)

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Light Grid Accent 1'
hdr = table4.rows[0].cells
hdr[0].text = 'Componente'
hdr[1].text = 'Costo mensual'

costos = [
    ('AWS S3 (backups)', '$10-20'),
    ('AWS KMS', '$1-3'),
    ('Neon Pro (PITR)', '$19'),
    ('GitHub Actions', 'Incluido'),
    ('Total estimado', '~$30-42/mes'),
]
for c in costos:
    row = table4.add_row().cells
    for i, val in enumerate(c):
        row[i].text = val

doc.add_paragraph()

# ============= PREGUNTAS =============
doc.add_heading('Preguntas de Validación', 1)
doc.add_paragraph(
    'Antes de proceder con la implementación, necesito clarificar:'
)
doc.add_paragraph('MFA: ¿TOTP solo, o también WebAuthn/FIDO2?', style='List Number')
doc.add_paragraph('MFA obligatorio: ¿Para todos los roles o solo superadmin y admin_empresa?')
doc.add_paragraph('Backup storage: ¿Tenés cuenta AWS/GCP, o preferís Vercel KV / otro?')
doc.add_paragraph('Cifrado: ¿Fernet (simétrico) o AWS KMS (asimétrico)?')
doc.add_paragraph('Compliance: ¿Hay auditoría externa de Ley 21.719 próximamente que requiera alguno de estos controles ya?')

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('— Fin del documento —').italic = True

# Save
output_path = 'messieufont/Plan_Mejoras_Seguridad_Custodio_RAT.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')