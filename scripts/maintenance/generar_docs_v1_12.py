"""
scripts/maintenance/generar_docs_v1_12.py
Genera los documentos oficiales v1.12 de Custodio RAT Manager.

Cambios respecto a v1.11:
- QW5 Formulario Público: detección de titular repetido (GET /publico/verificar-titular)
- ARCO-QW6: acuse de recibo automático al crear ticket
- ARCO-QW7: chips de placeholders en respuesta
- ARCO-QW8: banner SLA con tiempos reales en FlujoModal
- Empresas-QW6: CompanyFichaPanel (tabs Datos/RATs/ARCO/Brechas)
- CI/CD: pip-audit (CVE scan), env vars corregidas en backend-tests
- Nomenclatura APDP corregida en todo el código y tests

Ejecutar:
  cd backend
  .\\venv\\Scripts\\python.exe ..\\scripts\\maintenance\\generar_docs_v1_12.py
"""

import os
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERSION = "v1.12"
FECHA = date.today().strftime("%d/%m/%Y")
OUT_DIR = Path(__file__).parent.parent.parent / "docs" / "documentacion_oficial"


# ── Helpers ──────────────────────────────────────────────────────────────────

def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level, size, color in [
        ("Heading 1", 16, RGBColor(0x1F, 0x49, 0x7D)),
        ("Heading 2", 14, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 12, RGBColor(0x40, 0x40, 0x40)),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color


def add_cover(doc, titulo, subtitulo=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Custodio RAT Manager")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()
    h = doc.add_heading(titulo, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitulo:
        p2 = doc.add_paragraph(subtitulo)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Versión: {VERSION}  ·  Fecha: {FECHA}\nLey 21.719 — Protección de Datos Personales de Chile")
    doc.add_page_break()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    return table


# ── 08 API REST ───────────────────────────────────────────────────────────────

def generar_08_api():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "08 — API REST", "Referencia completa de endpoints")

    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(
        "Este documento describe todos los endpoints de la API REST de Custodio RAT Manager "
        "conforme a la Ley 21.719 de Chile. La API sigue el esquema /api/v1/ (canónico) "
        "manteniendo compatibilidad legacy. Versión v1.12 incluye el endpoint de verificación "
        "de titular repetido (Público-QW5) y correcciones de nomenclatura APDP."
    )

    ENDPOINTS_PUBLICO = [
        ("GET",  "/publico/empresas",           "—",   "—",         "—",                          "200 list[EmpresaPublica]",        "Lista empresas activas (id+nombre) para el formulario ARCO."),
        ("GET",  "/publico/csrf-token",         "—",   "—",         "—",                          "200 CsrfTokenResponse",           "Token HMAC-SHA256 para proteger formulario público (30/min)."),
        ("GET",  "/publico/verificar-titular",  "—",   "—",         "company_id, email",          "200 {tiene_tickets_abiertos, cantidad}", "Verifica si email ya tiene tickets abiertos. Rate 20/min. (QW5)"),
        ("POST", "/publico/ejercer-derechos",   "—",   "—",         "EjercerDerechosRequest",     "201 {tracking_token, mensaje}",   "Crea solicitud ARCO pública (10/hora por IP). Acuse email automático."),
        ("GET",  "/seguimiento/{token}",        "—",   "—",         "token (path)",               "200 SeguimientoOut",              "Consulta pública de estado de ticket por tracking token."),
        ("GET",  "/publico/transparencia/{id}", "—",   "—",         "company_id (path)",          "200 PoliticaTransparenciaOut",    "Política de transparencia pública (Art. 14 ter). Cache recomendado."),
    ]

    ENDPOINTS_AUTH = [
        ("POST", "/auth/login",    "—",   "—",          "username, password",     "200 {access_token}",         "Login JWT. Access token (8h) + refresh cookie httpOnly (30d)."),
        ("POST", "/auth/refresh",  "—",   "cookie",     "—",                      "200 {access_token}",         "Renovar access token usando refresh cookie."),
        ("POST", "/auth/logout",   "JWT", "cualquier",  "—",                      "204",                        "Revocar access + refresh tokens."),
        ("GET",  "/auth/me",       "JWT", "cualquier",  "—",                      "200 UserOut",                "Usuario autenticado actual."),
        ("POST", "/auth/users",    "JWT", "superadmin", "UserCreate",             "201 UserOut",                "Crear usuario (superadmin only)."),
    ]

    ENDPOINTS_ARCO = [
        ("GET",    "/tkt-solicitud-derecho/",                    "JWT", "admin_empresa+",  "filtros",              "200 TktListResponse",   "Listar tickets ARCO con filtros."),
        ("POST",   "/tkt-solicitud-derecho/",                    "JWT", "editor+",         "TktCreate",            "201 TktOut",            "Crear ticket ARCO. Envía acuse de recibo al titular (ARCO-QW6)."),
        ("GET",    "/tkt-solicitud-derecho/{id}",                "JWT", "cualquier",       "id (path)",            "200 TktOut",            "Detalle de ticket."),
        ("PATCH",  "/tkt-solicitud-derecho/{id}",                "JWT", "editor+",         "TktUpdate",            "200 TktOut",            "Actualizar ticket. Requiere metodo_verificacion_identidad al resolver."),
        ("POST",   "/tkt-solicitud-derecho/{id}/rechazar",       "JWT", "editor+",         "causal_rechazo",       "200 TktOut",            "Rechazar con causal (enum). Email al titular."),
        ("POST",   "/tkt-solicitud-derecho/{id}/prorrogar",      "JWT", "editor+",         "—",                    "200 TktOut",            "+10 días hábiles (Art. 12 bis). Una vez por ticket."),
        ("POST",   "/tkt-solicitud-derecho/{id}/bloquear",       "JWT", "editor+",         "plazo_bloqueo",        "200 TktOut",            "Bloquear RAT (Art. 8 ter)."),
        ("POST",   "/tkt-solicitud-derecho/{id}/subsanar",       "JWT", "editor+",         "detalle",              "200 TktOut",            "Solicitar subsanación al titular."),
        ("POST",   "/tkt-solicitud-derecho/{id}/completar-subsanacion", "JWT", "editor+", "—",                    "200 TktOut",            "Cerrar subsanación."),
        ("GET",    "/tkt-solicitud-derecho/{id}/portabilidad/export", "JWT", "editor+",   "id (path)",            "200 JSON",              "Exportar datos portabilidad (Art. 9)."),
    ]

    doc.add_heading("2. Endpoints Públicos (sin autenticación)", level=1)
    add_table(doc,
              ["Método", "Ruta", "Auth", "RBAC", "Parámetros", "Respuesta", "Descripción"],
              [(e[0], e[1], e[2], e[3], e[4], e[5], e[6]) for e in ENDPOINTS_PUBLICO])

    doc.add_heading("3. Autenticación", level=1)
    add_table(doc,
              ["Método", "Ruta", "Auth", "RBAC", "Parámetros", "Respuesta", "Descripción"],
              [(e[0], e[1], e[2], e[3], e[4], e[5], e[6]) for e in ENDPOINTS_AUTH])

    doc.add_heading("4. ARCO — Solicitudes de Derecho", level=1)
    doc.add_paragraph(
        "Tabla canónica: tkt_solicitud_derecho. La tabla legacy solicitudes_derecho fue eliminada en jul-2026."
    )
    add_table(doc,
              ["Método", "Ruta", "Auth", "RBAC", "Parámetros", "Respuesta", "Descripción"],
              [(e[0], e[1], e[2], e[3], e[4], e[5], e[6]) for e in ENDPOINTS_ARCO])

    doc.add_heading("5. Cambios en v1.12", level=1)
    cambios = [
        ("GET /publico/verificar-titular", "NUEVO", "QW5: detecta titular con tickets abiertos. Rate 20/min."),
        ("POST /tkt-solicitud-derecho/", "MODIFICADO", "ARCO-QW6: envía acuse de recibo email al crear ticket."),
        ("Nomenclatura APDP", "CORREGIDO", "APDC → APDP en todo el código, tests y documentación."),
        ("CI/CD pip-audit", "NUEVO", "Escaneo CVE en dependencias Python. Bloquea CRITICAL, advierte HIGH."),
    ]
    add_table(doc, ["Endpoint / Cambio", "Tipo", "Descripción"], cambios)

    path = OUT_DIR / f"08_API_REST_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── 04 Casos de Uso ───────────────────────────────────────────────────────────

def generar_04_casos_uso():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "04 — Casos de Uso", "Custodio RAT Manager · Ley 21.719")

    doc.add_heading("Resumen Ejecutivo", level=1)
    doc.add_paragraph(
        f"Este documento describe los casos de uso del sistema Custodio RAT Manager en versión {VERSION}. "
        "Cubre 30+ casos de uso incluyendo los nuevos módulos ARCO-QW5/6/7/8, Empresas-QW6 y "
        "la detección de titular repetido en el formulario público."
    )

    # Tabla de CUs nuevos en v1.12
    doc.add_heading("1. Casos de Uso Nuevos en v1.12", level=1)
    nuevos_cus = [
        ("CU-31", "Detección de titular repetido", "Formulario Público ARCO", "Anónimo (titular)",
         "El sistema verifica si el email del titular ya tiene tickets abiertos en la empresa seleccionada. "
         "Si existen, muestra banner amarillo de advertencia. El titular puede continuar de todos modos.",
         "Art. 12 Ley 21.719"),
        ("CU-32", "Acuse de recibo automático ARCO", "Módulo ARCO", "Staff (editor+)",
         "Al crear un ticket ARCO con email del titular, el sistema envía automáticamente un acuse "
         "de recibo con el tracking token y la fecha de vencimiento (10 días hábiles).",
         "Art. 12 Ley 21.719"),
        ("CU-33", "Insertar placeholders en respuesta", "Módulo ARCO", "Staff (editor+)",
         "El operador redacta la respuesta al titular. Chips de placeholders dinámicos permiten "
         "insertar {{nombre_titular}}, {{empresa}}, {{fecha}}, {{numero_solicitud}}, "
         "{{dias_bloqueo}}, {{fecha_vencimiento}} con un clic.",
         "Usabilidad"),
        ("CU-34", "Ver flujo ARCO con tiempos reales", "Módulo ARCO", "Staff (editor+)",
         "El modal de flujo muestra el estado actual del ticket junto con días hábiles consumidos, "
         "días restantes y un semáforo de color (verde/amarillo/rojo) según el SLA.",
         "Art. 12 Ley 21.719"),
        ("CU-35", "Ficha de empresa (tabs)", "Módulo Empresas", "Staff (cualquier rol)",
         "Al hacer clic en 'Ficha' de una empresa, se despliega un panel con 4 tabs: "
         "Datos (edición), RATs (listado), ARCO (tickets), Brechas. Carga lazy por tab.",
         "Art. 16 Ley 21.719"),
    ]
    add_table(doc,
              ["ID", "Nombre", "Módulo", "Actor", "Descripción", "Artículo"],
              nuevos_cus)

    doc.add_heading("2. Casos de Uso Existentes (referencia)", level=1)
    doc.add_paragraph(
        "Los siguientes CUs fueron documentados en versiones anteriores y permanecen vigentes. "
        "Ver 04_Casos_de_Uso_v1.10.docx para detalle completo."
    )
    cus_existentes = [
        ("CU-01 a CU-06", "Onboarding y Gestión de Empresas"),
        ("CU-07 a CU-14", "CRUD RAT (crear, editar, aprobar, duplicar, exportar)"),
        ("CU-15 a CU-18", "Brechas de Seguridad (Art. 14 bis)"),
        ("CU-19 a CU-22", "Consentimientos (Art. 12)"),
        ("CU-23 a CU-26", "EIPD (Art. 15 bis)"),
        ("CU-27 a CU-30", "Formulario Público ARCO y Seguimiento"),
    ]
    add_table(doc, ["Rango", "Descripción"], cus_existentes)

    doc.add_heading("3. Flujo Principal — CU-31: Detección Titular Repetido", level=1)
    doc.add_paragraph("Precondición: el titular está en el Paso 1 del formulario público con empresa ya seleccionada.")
    flujo = [
        ("1", "El titular ingresa su email en el campo correspondiente."),
        ("2", "Al salir del campo (onBlur), el frontend llama GET /publico/verificar-titular?company_id=X&email=Y."),
        ("3a", "Si {tiene_tickets_abiertos: false} → no ocurre nada, flujo normal."),
        ("3b", "Si {tiene_tickets_abiertos: true, cantidad: N} → aparece banner amarillo advirtiendo N ticket(s) abierto(s)."),
        ("4", "El titular puede continuar de todos modos (el banner no bloquea)."),
        ("5", "Si el titular decide continuar, avanza al Paso 2 normalmente."),
    ]
    add_table(doc, ["Paso", "Acción"], flujo)

    path = OUT_DIR / f"04_Casos_de_Uso_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── 09 Backlog ────────────────────────────────────────────────────────────────

def generar_09_backlog():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "09 — Backlog de Producto", f"Estado al {FECHA}")

    doc.add_heading("Estado Global", level=1)
    resumen = [
        ("Módulo Empresas", "5", "0", "5", "0", "10"),
        ("Módulo ARCO (Tickets)", "6", "0", "3", "1", "10"),
        ("Formulario Público ARCO", "4", "0", "6", "0", "10"),
        ("Formulario Admin ARCO", "0", "0", "10", "0", "10"),
        ("TOTAL", "15", "0", "24", "1", "40"),
    ]
    add_table(doc, ["Módulo", "Pendiente", "En Progreso", "Completado", "Postergado", "Total"], resumen)

    doc.add_heading("Módulo Empresas — Quick Wins", level=1)
    empresas_qws = [
        ("QW1", "Vista auditoría per-empresa", "ALTO", "✅ Cerrado", "2026-08-21"),
        ("QW2", "Exportar Reporte APDP (PDF)", "CRÍTICO", "✅ Cerrado", "2026-08-21"),
        ("QW3", "Score de cumplimiento v1", "MEDIO", "⚪ Pendiente", "—"),
        ("QW4", "Exportación CSV/Excel/PDF tickets", "ALTO", "⚪ Pendiente", "—"),
        ("QW5", "SLA alert email T-2 días", "CRÍTICO", "✅ Cerrado", "2026-07-18"),
        ("QW6", "Ficha empresa con tabs (Datos/RATs/ARCO/Brechas)", "MEDIO", "✅ Cerrado", "2026-08-22"),
        ("QW7", "Banner alertas en lista empresas", "MEDIO", "✅ Cerrado", "2026-08-21"),
        ("QW8", "Recordatorio ARCO T-2 días", "ALTO", "⚪ Pendiente", "—"),
        ("QW9", "Editar RUT post-creación", "BAJO", "⚪ Pendiente", "—"),
        ("QW10", "Plantillas RAT por rubro", "MEDIO", "⚪ Pendiente", "—"),
    ]
    add_table(doc, ["#", "Mejora", "Impacto", "Estado", "Fecha Cierre"], empresas_qws)

    doc.add_heading("Módulo ARCO — Quick Wins", level=1)
    arco_qws = [
        ("QW1", "Exportación CSV/Excel/PDF", "ALTO", "✅ Completado", "2026-06-24"),
        ("QW2", "SLA alert email T-2 días", "CRÍTICO", "✅ Completado", "2026-06-24"),
        ("QW3", "Firma digital + timestamp", "CRÍTICO", "⏸ Postergado", "—"),
        ("QW4", "Dashboard derechos más ejercidos", "BAJO", "✅ Cerrado", "2026-07-09"),
        ("QW5", "Bandeja de entrada DPO", "ALTO", "⚪ Pendiente", "—"),
        ("QW6", "Acuse de recibo automático al titular", "ALTO", "✅ Cerrado", "2026-08-22"),
        ("QW7", "Chips de placeholders en respuesta", "MEDIO", "✅ Cerrado", "2026-08-22"),
        ("QW8", "Ver Flujo con tiempos reales (banner SLA)", "MEDIO", "✅ Cerrado", "2026-08-22"),
        ("QW9", "Portal del titular con descarga", "ALTO", "⚪ Pendiente", "—"),
        ("QW10", "Editar RUT titular/representante", "BAJO", "⚪ Pendiente", "—"),
    ]
    add_table(doc, ["#", "Mejora", "Impacto", "Estado", "Fecha Cierre"], arco_qws)

    doc.add_heading("Formulario Público ARCO — Quick Wins", level=1)
    publico_qws = [
        ("QW1", "Validación RUT en vivo + formateo", "MEDIO", "✅ Cerrado", "2026-08-22"),
        ("QW2", "Banner privacidad + link política", "CRÍTICO", "✅ Completado", "2026-06-24"),
        ("QW3", "Pantalla intro + glosario", "ALTO", "✅ Cerrado", "2026-08-22"),
        ("QW4", "Confirmación de email (doble input)", "MEDIO", "✅ Cerrado", "2026-08-22"),
        ("QW5", "Detección de titular repetido", "MEDIO", "✅ Cerrado", "2026-08-24"),
        ("QW6", "CTA 'Consultar estado' en pantalla éxito", "MEDIO", "✅ Completado", "2026-06-24"),
        ("QW7", "Validación tamaño archivo", "BAJO", "⚪ Pendiente", "—"),
        ("QW8", "Eliminar archivos individuales upload", "BAJO", "⚪ Pendiente", "—"),
        ("QW9", "Stepper con 3 pasos visibles", "BAJO", "✅ Cerrado", "2026-08-22"),
        ("QW10", "Microcopy mejorado", "BAJO", "⚪ Pendiente", "—"),
    ]
    add_table(doc, ["#", "Mejora", "Impacto", "Estado", "Fecha Cierre"], publico_qws)

    path = OUT_DIR / f"09_Backlog_Producto_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── 12 Manual Técnico ─────────────────────────────────────────────────────────

def generar_12_manual_tecnico():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "12 — Manual Técnico", "Guía de arquitectura e implementación")

    doc.add_heading("1. Resumen de Cambios v1.12", level=1)
    cambios = [
        ("Público-QW5", "Backend + Frontend", "GET /publico/verificar-titular: detecta titular con tickets abiertos antes de enviar nueva solicitud ARCO."),
        ("ARCO-QW6", "Backend", "notificar_acuse_solicitud(): email automático al crear ticket ARCO con titular_email."),
        ("ARCO-QW7", "Frontend", "Chips de placeholders en TicketDrawer.tsx: {{nombre_titular}}, {{empresa}}, {{fecha}}, etc."),
        ("ARCO-QW8", "Frontend", "Banner SLA en FlujoModal.tsx: días hábiles consumidos, días restantes, semáforo de color."),
        ("Empresas-QW6", "Frontend", "CompanyFichaPanel.tsx: panel con 4 tabs (Datos/RATs/ARCO/Brechas) con carga lazy."),
        ("CI/CD", "DevOps", "pip-audit en job lint: bloquea CRITICAL, advierte HIGH. Env vars ALLOWED_ORIGINS y ENVIRONMENT hardcodeados para tests."),
        ("Nomenclatura", "Global", "APDC → APDP (Agencia de Protección de Datos Personales) en código, tests y docs."),
    ]
    add_table(doc, ["Componente", "Capa", "Descripción"], cambios)

    doc.add_heading("2. Nuevo Endpoint: GET /publico/verificar-titular", level=1)
    doc.add_paragraph("Archivo: backend/app/routes/publico_arco.py")
    doc.add_paragraph(
        "Endpoint público (sin autenticación) que verifica si un email ya tiene solicitudes ARCO "
        "en estado abierto/en_proceso/pendiente_subsanacion para una empresa dada.\n\n"
        "Rate limit: 20/minuto por IP (para evitar enumeración de emails).\n"
        "Respuesta: {tiene_tickets_abiertos: bool, cantidad: int}"
    )
    doc.add_heading("2.1 Seguridad", level=2)
    doc.add_paragraph(
        "• No revela si el email existe en el sistema (solo indica tickets abiertos en esa empresa).\n"
        "• Rate limited con slowapi (20/min) para prevenir ataques de enumeración.\n"
        "• Solo estados activos: abierto, en_proceso, pendiente_subsanacion, pendiente.\n"
        "• titular_email almacenado como texto plano (String(255)) — no cifrado en tkt_solicitud_derecho."
    )

    doc.add_heading("3. ARCO-QW6: Acuse de Recibo Automático", level=1)
    doc.add_paragraph("Archivo: backend/app/routes/tkt_solicitud_derecho.py")
    doc.add_paragraph(
        "Al crear un ticket ARCO (POST /tkt-solicitud-derecho/) con titular_email presente, "
        "el sistema llama notificar_acuse_solicitud() del email_service. Si falla, registra "
        "warning en logs pero no revierte la creación del ticket (no bloquea el flujo)."
    )

    doc.add_heading("4. Empresas-QW6: CompanyFichaPanel", level=1)
    doc.add_paragraph("Archivo: frontend-next/components/companies/CompanyFichaPanel.tsx")
    doc.add_paragraph(
        "Componente de panel con 4 tabs con carga lazy:\n"
        "• Datos: edición de la empresa (reutiliza CompanyEditForm)\n"
        "• RATs: listado de RATs de la empresa con completitud\n"
        "• ARCO: tickets de solicitudes de derecho (usa TktListResponse.tickets)\n"
        "• Brechas: listado de brechas de seguridad\n\n"
        "Se activa con un botón 'Ficha' por tarjeta de empresa en /companies."
    )

    doc.add_heading("5. CI/CD — Pipeline Tests + Coverage", level=1)
    doc.add_paragraph("Archivo: .github/workflows/tests.yml")
    ci_cambios = [
        ("pip-audit", "lint", "Escaneo CVE en requirements.txt. CRITICAL bloquea CI; HIGH aparece como warning."),
        ("ALLOWED_ORIGINS", "backend-tests", "Hardcodeado a 'http://localhost:3000' (evita RuntimeError al importar main.py)."),
        ("ENVIRONMENT", "backend-tests", "Hardcodeado a 'development' (activa dev fallbacks de crypto y JWT)."),
        ("Verificar secret", "backend-tests", "Step previo que falla con mensaje claro si TEST_DATABASE_URL no está configurado."),
        ("include/exclude vitest", "frontend-tests", "Excluye e2e/ para que Playwright no corra dentro de vitest."),
    ]
    add_table(doc, ["Cambio", "Job", "Detalle"], ci_cambios)

    doc.add_heading("6. Stack Tecnológico (vigente)", level=1)
    stack = [
        ("Backend", "FastAPI 0.115 + Uvicorn", "Python 3.9+"),
        ("ORM", "SQLAlchemy 2.0", "PostgreSQL (Neon)"),
        ("Validación", "Pydantic 2.10", "—"),
        ("Auth", "JWT (HS256) + bcrypt", "python-jose 3.3.0, passlib 1.7.4"),
        ("Cifrado", "Fernet (cryptography 43)", "ENCRYPTION_KEY obligatoria en prod/qa"),
        ("Email", "SMTP via SMTP_URL DSN", "Modo DRY_RUN si no configurado"),
        ("Frontend", "Next.js 16.2 + React 19", "TypeScript + Tailwind CSS v4"),
        ("Tests backend", "pytest 8.3 + httpx", "761+ tests contra Neon QA"),
        ("Tests frontend", "vitest 3 + Playwright", "E2E separados del runner vitest"),
        ("CI/CD", "GitHub Actions", "3 jobs: backend-tests, frontend-tests, lint (ruff + eslint + pip-audit)"),
        ("Hosting", "Vercel", "Node.js 24 (actualizado desde 20)"),
        ("Base de datos", "Neon PostgreSQL", "BD prod + BD QA separada para tests"),
    ]
    add_table(doc, ["Componente", "Tecnología", "Notas"], stack)

    path = OUT_DIR / f"12_Manual_Tecnico_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── 10 Plan QA ────────────────────────────────────────────────────────────────

def generar_10_plan_qa():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "10 — Plan de QA", f"Suite de pruebas · {FECHA}")

    doc.add_heading("1. Métricas Actuales", level=1)
    metricas = [
        ("Tests backend (pytest)", "761+", "Neon QA (custodio_test)"),
        ("Tests frontend (vitest)", "Incluidos en CI", "vitest 3, excluye e2e/"),
        ("Tests E2E (Playwright)", "Ejecutados por separado", "playwright.config.ts"),
        ("Cobertura backend", "Reportada via codecov", "coverage.xml"),
        ("Cobertura frontend", "Reportada via codecov", "lcov.info"),
        ("Fallos en suite", "0", "Post-QA total 2026-08-22 (78 → 0)"),
    ]
    add_table(doc, ["Métrica", "Valor", "Notas"], metricas)

    doc.add_heading("2. Nuevos Casos de Prueba v1.12", level=1)
    nuevos_tc = [
        ("TC-056", "QW5 titular nuevo", "GET /publico/verificar-titular con email sin tickets → tiene_tickets_abiertos=false"),
        ("TC-057", "QW5 titular repetido", "GET /publico/verificar-titular con email con ticket abierto → tiene_tickets_abiertos=true, cantidad>=1"),
        ("TC-058", "QW5 titular cerrado", "GET /publico/verificar-titular con email solo con ticket resuelto → tiene_tickets_abiertos=false"),
        ("TC-059", "QW5 rate limit", "21 llamadas en 1 minuto desde misma IP → respuesta 429"),
        ("TC-060", "ARCO-QW6 acuse email", "POST ticket con titular_email → acuse_enviado_at se popula en BD"),
        ("TC-061", "ARCO-QW6 sin email", "POST ticket sin titular_email → no error, acuse_enviado_at=null"),
        ("TC-062", "Empresas-QW6 tab RATs", "GET /rats/?company_id=X desde CompanyFichaPanel → retorna RATs de la empresa"),
        ("TC-063", "APDP nomenclatura", "GET /rats/export/cni → título contiene 'APDP' (no 'APDC')"),
    ]
    add_table(doc, ["TC", "Nombre", "Descripción / Expectativa"], nuevos_tc)

    doc.add_heading("3. Configuración CI/CD", level=1)
    doc.add_paragraph(
        "El pipeline 'Tests + Coverage' corre en GitHub Actions en cada push a main/qa/develop.\n\n"
        "Variables de entorno requeridas en GitHub Secrets:\n"
        "• TEST_DATABASE_URL — connection string a Neon QA (custodio_test)\n\n"
        "Variables hardcodeadas en el workflow (no sensibles):\n"
        "• ALLOWED_ORIGINS=http://localhost:3000\n"
        "• ENVIRONMENT=development\n\n"
        "El job 'lint' incluye desde v1.12:\n"
        "• ruff check backend/app/ (falla en E/F errors)\n"
        "• pip-audit -r backend/requirements.txt --severity critical (bloquea CRITICAL)\n"
        "• npm run lint (ESLint frontend)"
    )

    path = OUT_DIR / f"10_Plan_QA_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── 02 Requisitos ─────────────────────────────────────────────────────────────

def generar_02_requisitos():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "02 — Requisitos", "Requisitos Funcionales y No Funcionales")

    doc.add_heading("1. Nuevos Requisitos v1.12", level=1)

    doc.add_heading("1.1 Requisitos Funcionales", level=2)
    nuevos_rfs = [
        ("RF-174", "Verificar titular repetido (público)", "MEDIO",
         "El sistema debe verificar si un email ya tiene tickets ARCO abiertos en una empresa "
         "antes de que el titular envíe una nueva solicitud. La verificación es opcional (no bloquea "
         "el envío) y se realiza vía GET /publico/verificar-titular."),
        ("RF-175", "Acuse de recibo automático ARCO", "ALTO",
         "Al crear un ticket ARCO con email del titular, el sistema debe enviar un acuse de recibo "
         "automático con tracking token, empresa, tipo de derecho y fecha de vencimiento. "
         "Si el envío falla, se registra warning sin revertir el ticket."),
        ("RF-176", "Placeholders dinámicos en respuesta ARCO", "MEDIO",
         "El operador debe poder insertar variables dinámicas en el texto de respuesta mediante chips "
         "de selección: {{nombre_titular}}, {{empresa}}, {{fecha}}, {{numero_solicitud}}, "
         "{{dias_bloqueo}}, {{fecha_vencimiento}}."),
        ("RF-177", "Banner SLA en flujo ARCO", "MEDIO",
         "El modal de flujo debe mostrar los días hábiles consumidos y restantes del SLA legal "
         "con un semáforo de color: verde (>5 días), amarillo (2-5 días), rojo (<2 días)."),
        ("RF-178", "Ficha de empresa con tabs", "MEDIO",
         "El listado de empresas debe ofrecer acceso a una ficha de empresa con tabs: "
         "Datos (edición), RATs (listado), ARCO (tickets), Brechas. Cada tab carga datos al seleccionarse."),
        ("RF-179", "Escaneo de dependencias CVE", "ALTO",
         "El pipeline CI debe incluir escaneo automático de vulnerabilidades CVE en las dependencias "
         "Python. Las vulnerabilidades CRITICAL bloquean el deploy; las HIGH generan advertencias."),
    ]
    add_table(doc, ["ID", "Nombre", "Prioridad", "Descripción"], nuevos_rfs)

    doc.add_heading("1.2 Requisitos No Funcionales", level=2)
    nuevos_rnfs = [
        ("RNF-21", "Nomenclatura APDP", "Toda referencia a la agencia reguladora chilena debe usar 'APDP' "
         "(Agencia de Protección de Datos Personales), no 'APDC'."),
        ("RNF-22", "Rate limit verificación titular", "GET /publico/verificar-titular debe estar limitado "
         "a 20 llamadas/minuto por IP para prevenir enumeración de emails."),
    ]
    add_table(doc, ["ID", "Nombre", "Descripción"], nuevos_rnfs)

    doc.add_heading("2. Requisitos Anteriores Vigentes", level=1)
    doc.add_paragraph(
        "Los requisitos RF-001 a RF-173 documentados en versiones anteriores permanecen vigentes. "
        "Ver 02_Requisitos_v1.10.docx para el detalle completo."
    )

    path = OUT_DIR / f"02_Requisitos_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── 03 Historias de Usuario ───────────────────────────────────────────────────

def generar_03_historias():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "03 — Historias de Usuario", "User Stories · Ley 21.719")

    doc.add_heading("1. Nuevas Historias de Usuario v1.12", level=1)
    nuevas_hus = [
        ("HU-104", "Titular", "Detección solicitud duplicada",
         "Como titular, quiero saber si ya tengo una solicitud abierta antes de enviar otra, "
         "para no duplicar trámites innecesariamente.",
         "Banner amarillo al ingresar email con ticket abierto. No bloquea el envío.", "RF-174"),
        ("HU-105", "Titular", "Acuse de recibo inmediato",
         "Como titular, quiero recibir un email de confirmación al enviar mi solicitud ARCO, "
         "con el código de seguimiento y el plazo de respuesta.",
         "Email enviado automáticamente dentro de los 5 minutos de creado el ticket.", "RF-175"),
        ("HU-106", "Operador DPO", "Placeholders en respuesta",
         "Como operador DPO, quiero insertar variables dinámicas en la respuesta al titular "
         "con un clic, para redactar más rápido y con menos errores.",
         "Chips de placeholders disponibles bajo el textarea de respuesta.", "RF-176"),
        ("HU-107", "Operador DPO", "Ver tiempos reales del ticket",
         "Como operador DPO, quiero ver cuántos días hábiles lleva el ticket y cuántos quedan, "
         "para priorizar correctamente.",
         "Banner con semáforo verde/amarillo/rojo en el modal de flujo.", "RF-177"),
        ("HU-108", "Admin empresa", "Ficha completa de empresa",
         "Como administrador, quiero ver una ficha de empresa con todos sus módulos (RATs, ARCO, "
         "Brechas) en un solo lugar, para tener una visión consolidada.",
         "Panel con 4 tabs de carga lazy en la página de empresas.", "RF-178"),
    ]
    add_table(doc,
              ["ID", "Actor", "Nombre", "Historia", "Criterio de aceptación", "RF"],
              nuevas_hus)

    doc.add_heading("2. Historias Anteriores Vigentes", level=1)
    doc.add_paragraph("Las HU-001 a HU-103 documentadas en versiones anteriores permanecen vigentes.")

    path = OUT_DIR / f"03_Historias_Usuario_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── MTX Trazabilidad ──────────────────────────────────────────────────────────

def generar_mtx():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, "MTX — Matriz de Trazabilidad", f"RF ↔ HU ↔ CU ↔ Endpoint · {VERSION}")

    doc.add_heading("1. Trazabilidad Nuevos Elementos v1.12", level=1)
    mtx = [
        ("RF-174", "HU-104", "CU-31", "GET /publico/verificar-titular", "Art. 12 Ley 21.719", "TC-056, TC-057, TC-058, TC-059"),
        ("RF-175", "HU-105", "CU-32", "POST /tkt-solicitud-derecho/ (acuse)", "Art. 12 Ley 21.719", "TC-060, TC-061"),
        ("RF-176", "HU-106", "CU-33", "Frontend TicketDrawer.tsx (chips)", "Usabilidad", "—"),
        ("RF-177", "HU-107", "CU-34", "Frontend FlujoModal.tsx (banner SLA)", "Art. 12 Ley 21.719", "—"),
        ("RF-178", "HU-108", "CU-35", "Frontend CompanyFichaPanel.tsx", "Art. 16 Ley 21.719", "TC-062"),
        ("RF-179", "—", "—", "CI/CD pip-audit", "Seguridad DevOps", "—"),
        ("RNF-21", "—", "—", "Global (código, tests, docs)", "Nomenclatura APDP", "TC-063"),
        ("RNF-22", "—", "—", "GET /publico/verificar-titular", "Seguridad rate limit", "TC-059"),
    ]
    add_table(doc, ["RF", "HU", "CU", "Implementación", "Artículo Legal", "Tests"], mtx)

    doc.add_heading("2. Trazabilidad Anterior Vigente", level=1)
    doc.add_paragraph(
        "La trazabilidad de RF-001 a RF-173 / HU-001 a HU-103 / CU-01 a CU-30 permanece vigente. "
        "Ver Matriz_Trazabilidad_v1.10.docx para el detalle completo."
    )

    path = OUT_DIR / f"Matriz_Trazabilidad_Custodio_RAT_Manager_{VERSION}.docx"
    doc.save(str(path))
    print(f"✓ {path.name}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Generando documentos {VERSION} en {OUT_DIR}/\n")
    generar_08_api()
    generar_04_casos_uso()
    generar_09_backlog()
    generar_12_manual_tecnico()
    generar_10_plan_qa()
    generar_02_requisitos()
    generar_03_historias()
    generar_mtx()
    print(f"\n✓ 8 documentos generados en {VERSION}")
