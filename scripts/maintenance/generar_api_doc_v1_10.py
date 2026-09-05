"""
scripts/maintenance/generar_api_doc_v1_10.py
Genera 08_API_REST_Custodio_RAT_Manager_v1.10.docx con los 20 endpoints RAT
(antes v1.9 tenia solo 6 endpoints documentados — ver auditoria 2026-07-07).

Ejecutar:
  cd backend && python ../scripts/maintenance/generar_api_doc_v1_10.py
"""

import os
import sys
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Catalogo de endpoints (fuente: backend/app/routes/rats.py + auditoria 2026-07-07)
ENDPOINTS = [
    # (method, path, auth, rbac, params, response, tags, description)
    ("GET",  "/rats/reportes",                       "JWT", "cualquier rol con acceso",  "skip, limit, sort_by, sort_order, search, estado, base_legal, categoria_titulares, datos_sensibles, evaluacion_impacto, transferencia_internacional, created_by, categoria_datos, datos_nna, transferencia_nacional, nivel_confidencialidad, decisiones_automatizadas, company_id",  "200 + ReportesResponse",  "Reportes",       "Reporte filtrado con 18+ filtros, paginacion y ordenamiento (QW-ITER14-01)"),
    ("GET",  "/rats/",                                "JWT", "cualquier rol con acceso",  "company_id, skip, limit",                                                    "200 + list[RATOut]",      "RAT CRUD",       "Listar RATs. Multi-tenant. Feature gate N-02."),
    ("GET",  "/rats/dashboard/{company_id}",          "JWT", "admin_empresa, superadmin", "company_id (path)",                                                          "200 + DashboardStats",    "RAT Dashboard",  "KPIs: total, por estado, sensibles, EIPD, transferencias, IL sin test, encargados sin contrato, sin doc."),
    ("GET",  "/rats/sugerencias/tipos",               "JWT", "cualquier rol",             "-",                                                                          "200 + {tipos: list}",    "Sugerencias",    "Lista de tipos de proceso disponibles para sugerencias automaticas."),
    ("POST", "/rats/sugerencias",                     "JWT", "cualquier rol",             "tipo_proceso",                                                               "200 + RATSugerenciaOut", "Sugerencias",    "Dado un tipo de proceso, retorna sugerencias precompletadas para el RAT."),
    ("GET",  "/rats/{rat_id}",                        "JWT", "admin_empresa, usuario",    "rat_id (path)",                                                              "200 + RATOut",           "RAT CRUD",       "Obtener RAT por ID. IDOR prevention: get_rat_for_user retorna 404 si no pertenece."),
    ("POST", "/rats/",                                "JWT", "editor, admin_empresa",     "RATCreate (40 campos)",                                                      "201 + RATOut",           "RAT CRUD",       "Crear RAT. Valida EIPD obligatoria, consentimientos, contratos."),
    ("POST", "/rats/{rat_id}/consentimientos",         "JWT", "editor, admin_empresa",     "rat_id (path), ConsentimientoCreate",                                        "201 + ConsentimientoOut","Consentimiento", "Registrar consentimiento expreso (Art. 12). Cifrado PII con Fernet + SHA-256."),
    ("PUT",  "/rats/{rat_id}",                        "JWT", "editor, admin_empresa",     "rat_id (path), RATUpdate",                                                   "200 + RATOut",           "RAT CRUD",       "Actualizar RAT. Validadores condicionales (transferencia_int, decisiones_auto, datos_sensibles)."),
    ("DELETE","/rats/{rat_id}",                       "JWT", "editor, admin_empresa",     "rat_id (path)",                                                              "200 + {message}",        "RAT CRUD",       "Eliminar RAT. Mueve archivo a archive bucket antes de borrar."),
    ("POST", "/rats/{rat_id}/revision",               "JWT", "editor, admin_empresa",     "rat_id (path)",                                                              "200 + AuditLogOut",      "RAT Lifecycle",  "Marcar el proceso como revisado periodicamente."),
    ("POST", "/rats/{rat_id}/aprobar",                "JWT", "admin_empresa, superadmin", "rat_id (path)",                                                              "200 + RATOut",           "RAT Lifecycle",  "Aprobar un RAT. Requiere 100% completitud."),
    ("GET",  "/rats/{rat_id}/archivo",                "JWT", "editor, admin_empresa",     "rat_id (path)",                                                              "200 + bytes | presigned", "RAT File",       "Descargar documento de base legal. Cadena: OCI → BYTEA descifrado Fernet."),
    ("GET",  "/rats/{rat_id}/auditoria",              "JWT", "cualquier rol con acceso",  "rat_id (path)",                                                              "200 + list[AuditLogOut]", "Auditoria",      "Historial de auditoria del RAT. IDOR prevention: get_rat_for_user."),
    ("GET",  "/rats/auditoria/{company_id}",          "JWT", "admin_empresa, superadmin", "company_id (path), skip, limit",                                             "200 + list[AuditLog]",   "Auditoria",      "Auditoria global de la empresa. Filtra por rat_ids de la empresa."),
    ("GET",  "/rats/auditoria/verify-chain",          "JWT", "solo SUPERADMIN",           "limit",                                                                      "200 + {valido, error_id}", "Auditoria",    "Verificar integridad de la cadena de hashes SHA-256 (H2.2 — restringido a SUPERADMIN)."),
    ("GET",  "/rats/export/csv",                      "JWT", "admin_empresa, superadmin", "company_id (query)",                                                          "200 + CSV text/csv",     "Export",         "Exportar todos los RATs de la empresa a CSV (UTF-8 BOM, sanitizado contra injection)."),
    ("GET",  "/rats/export/pdf",                      "JWT", "admin_empresa, superadmin", "company_id (query)",                                                          "200 + PDF",              "Export",         "Exportar todos los RATs a PDF. Marca 'RAT BLOQUEADO' en rojo (Art. 8 ter)."),
    ("GET",  "/rats/{rat_id}/export/pdf",             "JWT", "admin_empresa, superadmin", "rat_id (path)",                                                              "200 + PDF",              "Export",         "Exportar un RAT individual a PDF."),
    ("GET",  "/rats/export/cni",                      "JWT", "admin_empresa, superadmin", "company_id (query)",                                                          "200 + text/plain",       "Export",         "Exportar RAT en formato APDC (Ley 21.719)."),
]


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for level, size, color in [
        ('Heading 1', 16, RGBColor(0x1F, 0x49, 0x7D)),
        ('Heading 2', 14, RGBColor(0x2E, 0x74, 0xB5)),
        ('Heading 3', 12, RGBColor(0x40, 0x40, 0x40)),
    ]:
        s = doc.styles[level]
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color


def add_metadata_table(doc, version, fecha, autor, cambios):
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid'
    headers = ['Version', 'Fecha', 'Autor', 'Cambios']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    row1 = table.rows[1].cells
    row1[0].text = version
    row1[1].text = fecha
    row1[2].text = autor
    row1[3].text = cambios
    doc.add_paragraph()


def add_endpoints_table(doc, endpoints):
    """Tabla principal con los 20 endpoints."""
    headers = ['Metodo', 'Path', 'Auth', 'RBAC', 'Params', 'Response', 'Tags', 'Descripcion']
    table = doc.add_table(rows=1 + len(endpoints), cols=len(headers))
    table.style = 'Light Grid'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    for r, (method, path, auth, rbac, params, response, tags, desc) in enumerate(endpoints, start=1):
        row = table.rows[r].cells
        row[0].text = method
        row[1].text = path
        row[2].text = auth
        row[3].text = rbac
        row[4].text = params
        row[5].text = response
        row[6].text = tags
        row[7].text = desc


def add_endpoint_detail(doc, method, path, auth, rbac, params, response, tags, desc):
    h = doc.add_heading(f"{method} {path}", level=2)
    p = doc.add_paragraph()
    p.add_run("Auth: ").bold = True
    p.add_run(auth)
    p.add_run("    RBAC: ").bold = True
    p.add_run(rbac)
    p.add_run("    Tags: ").bold = True
    p.add_run(tags)
    if params and params != "-":
        doc.add_paragraph().add_run("Parametros: ").bold = True
        doc.add_paragraph(params, style='Intense Quote')
    p = doc.add_paragraph()
    p.add_run("Response: ").bold = True
    p.add_run(response)
    p = doc.add_paragraph()
    p.add_run("Descripcion: ").bold = True
    p.add_run(desc)


def main():
    # Output path
    script_dir = Path(__file__).resolve().parent.parent.parent  # scripts/maintenance -> RAT_opencode
    docs_dir = script_dir / "docs" / "documentacion_oficial"
    output = docs_dir / "08_API_REST_Custodio_RAT_Manager_v1.10.docx"

    doc = Document()
    setup_styles(doc)

    # Title
    doc.add_heading('CUSTODIO RAT — API REST v1.10', 0)
    doc.add_paragraph('Sistema RAT Manager — Ley 21.719 de Proteccion de Datos Personales de Chile')
    doc.add_paragraph()

    # Metadata
    doc.add_heading('Metadata', 1)
    add_metadata_table(
        doc,
        version='1.10',
        fecha='2026-07-07',
        autor='Auditoria RAT detallada — Custodio',
        cambios='Regeneracion con 20 endpoints RAT (antes 6 en v1.9). Auditoria 2026-07-07.',
    )

    # Resumen
    doc.add_heading('Resumen', 1)
    doc.add_paragraph(
        'Este documento describe los 20 endpoints REST del modulo RAT (Registro de Actividades de Tratamiento) '
        'segun la Ley 21.719 de Chile. Los endpoints cumplen con:'
    )
    bullets = [
        'Autenticacion JWT en todos los endpoints (excepto /publico/*).',
        'Multi-tenant defense via get_rat_for_user() con patron 404 (no exponer existencia).',
        'RBAC granular: SUPERADMIN, ADMIN_EMPRESA, USUARIO con jerarquia explicita.',
        'Audit log con cadena de hashes SHA-256 para integridad.',
        'CSV injection prevention con _DANGEROUS_CSV_PREFIXES.',
        'Cifrado PII (Fernet) para consentimientos (Art. 11 deber de confidencialidad).',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Tabla general de endpoints
    doc.add_heading(f'Tabla de endpoints ({len(ENDPOINTS)} total)', 1)
    doc.add_paragraph('A continuacion los 20 endpoints del modulo RAT:')
    add_endpoints_table(doc, ENDPOINTS)

    doc.add_page_break()

    # Detalle de cada endpoint
    doc.add_heading('Detalle de endpoints', 1)
    for ep in ENDPOINTS:
        add_endpoint_detail(doc, *ep)

    # Compliance
    doc.add_page_break()
    doc.add_heading('Compliance Ley 21.719', 1)
    doc.add_paragraph('Los endpoints implementan los siguientes articulos de la Ley 21.719:')
    arts = [
        ('Art. 12', 'Consentimiento', '/rats/{rat_id}/consentimientos — registro con cifrado PII'),
        ('Art. 14 bis', 'Brechas', 'Módulo separado (no parte del RAT directo)'),
        ('Art. 14 quater', 'Encargados', 'Validacion cruzada RAT↔Contrato en POST/PUT /rats/'),
        ('Art. 14 ter', 'Transparencia', 'Módulo separado (no parte del RAT directo)'),
        ('Art. 15 bis', 'EIPD', 'POST /rats/{id}/aprobar requiere EIPD completada'),
        ('Art. 16', 'RAT', 'GET/POST/PUT/DELETE /rats/ + reportes'),
        ('Art. 16 BIS', 'Biometricos', 'Base legal especifica en POST /rats/'),
        ('Art. 8 ter', 'Bloqueo', 'Flag bloqueado en exports PDF/CSV'),
    ]
    table = doc.add_table(rows=1 + len(arts), cols=3)
    table.style = 'Light Grid'
    headers = ['Articulo', 'Tema', 'Endpoint / Comportamiento']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for r, (art, tema, endpoint) in enumerate(arts, start=1):
        row = table.rows[r].cells
        row[0].text = art
        row[1].text = tema
        row[2].text = endpoint

    # Multi-tenant security
    doc.add_page_break()
    doc.add_heading('Seguridad Multi-Tenant', 1)
    doc.add_paragraph(
        'Todos los endpoints de recurso especifico (/rats/{rat_id}/*) usan get_rat_for_user() que valida:'
    )
    items = [
        'RAT existe (404 si no existe).',
        'Usuario pertenece a la empresa del RAT (404 si no — no exponer existencia).',
        'RBAC: editor o admin_empresa para escritura.',
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_paragraph(
        'Patron 404 (no 403) — decision de seguridad para no filtrar la existencia del RAT a '
        'usuarios no autorizados. El test rat_auditoria_test.py::test_auditoria_idor_usuario_ajeno_404 '
        'verifica este comportamiento.'
    )

    # Auditoria 2026-07-07
    doc.add_heading('Cambios desde v1.9', 1)
    doc.add_paragraph(
        'Auditoria detallada del RAT (2026-07-07) detecto que v1.9 solo documentaba 6 endpoints. '
        'Esta v1.10 documenta los 20 endpoints reales del codigo en backend/app/routes/rats.py.'
    )
    cambios = [
        'Agregados: GET /rats/reportes, /rats/dashboard/{id}, /sugerencias/tipos, POST /sugerencias, '
        'POST /rats/{id}/consentimientos, GET /rats/{id}/archivo, GET /rats/auditoria/{id}, '
        'GET /rats/auditoria/verify-chain, GET /rats/export/csv, /export/pdf, /{id}/export/pdf, /export/cni.',
        'H2.2 (P1 — auditoria 2026-07-07): /rats/auditoria/verify-chain restringido a SUPERADMIN.',
        'H1.1 (P1): base_legal="Otra" requiere archivo adjunto (Art. 11+16 Ley 21.719).',
        'Endpoints de export documentados con su sanitizacion CSV y alerta de bloqueo PDF.',
    ]
    for c in cambios:
        doc.add_paragraph(c, style='List Bullet')

    # Save
    docs_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"OK: Generado {output}")
    print(f"   Endpoints: {len(ENDPOINTS)}")
    print(f"   Bytes: {output.stat().st_size}")


if __name__ == "__main__":
    main()