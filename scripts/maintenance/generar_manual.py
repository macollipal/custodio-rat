from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(14)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Calibri'
h3_style.font.size = Pt(12)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def add_shaded_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    return p

doc.add_heading('CUSTODIO - Manual de Usuario', 0)
doc.add_paragraph('Sistema RAT Manager · Ley 21.719 de Protección de Datos Personales de Chile')
doc.add_paragraph()

doc.add_paragraph('ÍNDICE')
doc.add_paragraph('1. Iniciar y detener el sistema', style='List Number')
doc.add_paragraph('2. Iniciar sesión', style='List Number')
doc.add_paragraph('3. Gestión de empresas', style='List Number')
doc.add_paragraph('4. Gestión de usuarios y accesos', style='List Number')
doc.add_paragraph('5. Dashboard - Vista general', style='List Number')
doc.add_paragraph('6. Gestión de procesos RAT', style='List Number')
doc.add_paragraph('7. Gestión de brechas de seguridad', style='List Number')
doc.add_paragraph('8. Reportes y exportaciones', style='List Number')
doc.add_paragraph('9. Configuración de cuenta', style='List Number')

doc.add_page_break()

doc.add_heading('1. INICIAR Y DETENER EL SISTEMA', 1)

doc.add_heading('1.1 Iniciar el sistema', 2)
p = doc.add_paragraph()
p.add_run('Ubicación del archivo: ').bold = True
p.add_run('Carpeta raíz del proyecto')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Hacer doble clic en el archivo iniciar.bat')
doc.add_paragraph('2. Se abrirá una ventana de comandos mostrando el progreso:')
doc.add_paragraph('   - [1/3] Verificando configuraciones...')
doc.add_paragraph('   - [2/3] Iniciando Backend (FastAPI)...')
doc.add_paragraph('   - [3/3] Iniciando Frontend (Next.js)...')
doc.add_paragraph('3. El navegador se abrirá automáticamente en http://localhost:3000')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('URLs de acceso:').bold = True
doc.add_paragraph('   - Frontend: http://localhost:3000')
doc.add_paragraph('   - Backend/Documentación API: http://localhost:8002/docs')

doc.add_heading('1.2 Detener el sistema', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Hacer doble clic en el archivo detener.bat')
doc.add_paragraph('2. El script cerrará las ventanas del Backend y Frontend')
doc.add_paragraph('3. Mostrará confirmación: "Sistema detenido correctamente"')

doc.add_page_break()

doc.add_heading('2. INICIAR SESIÓN', 1)

doc.add_heading('2.1 Pantalla de login', 2)
doc.add_paragraph('Acceso: http://localhost:3000/login')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Campos del formulario:').bold = True
doc.add_paragraph('   - Usuario: campo de texto')
doc.add_paragraph('   - Contraseña: campo de contraseña')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Ingresar nombre de usuario')
doc.add_paragraph('2. Ingresar contraseña')
doc.add_paragraph('3. Hacer clic en "Iniciar sesión"')
doc.add_paragraph('4. Si las credenciales son correctas, se redirige al Dashboard')

doc.add_heading('2.2 Cerrar sesión', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En el Sidebar izquierdo, hacer clic en "Cerrar sesión"')
doc.add_paragraph('2. Se limpian los datos de sesión y se redirige a /login')

doc.add_page_break()

doc.add_heading('3. GESTIÓN DE EMPRESAS', 1)
doc.add_paragraph('Acceso: Menú lateral → "Empresas" o botón en Dashboard')

doc.add_heading('3.1 Ver lista de empresas', 2)
doc.add_paragraph('Se muestra una tarjeta por cada empresa con:')
doc.add_paragraph('   - Razón social y RUT')
doc.add_paragraph('   - Rubro/Sector')
doc.add_paragraph('   - Nombre y email del DPO')
doc.add_paragraph('   - Cantidad de procesos RAT')
doc.add_paragraph('   - Badge "ACTIVA" en la empresa seleccionada')

doc.add_heading('3.2 Crear nueva empresa', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Hacer clic en "+ Nueva empresa" (header)')
doc.add_paragraph('2. Completar el formulario:')
doc.add_paragraph('   - Razón social * (obligatorio)')
doc.add_paragraph('   - RUT * (obligatorio, con validación automática)')
doc.add_paragraph('   - Rubro/Sector')
doc.add_paragraph('   - Dirección')
doc.add_paragraph('   - Nombre del DPO')
doc.add_paragraph('   - Email del DPO')
doc.add_paragraph('   - Descripción')
doc.add_paragraph('3. Hacer clic en "Crear empresa"')
doc.add_paragraph('4. La nueva empresa se crea y se selecciona automáticamente')

doc.add_heading('3.3 Editar empresa', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En la tarjeta de la empresa, hacer clic en "Editar"')
doc.add_paragraph('2. Se expande un formulario editable (excepto RUT)')
doc.add_paragraph('3. Modificar los campos deseados')
doc.add_paragraph('4. Hacer clic en "Guardar cambios"')

doc.add_heading('3.4 Eliminar empresa', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En la tarjeta de la empresa, hacer clic en "Eliminar"')
doc.add_paragraph('2. Confirmar en el diálogo: "¿Está seguro de eliminar?"')
doc.add_paragraph('3. La empresa se elimina. Si era activa, se selecciona otra.')

doc.add_heading('3.5 Seleccionar empresa activa', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Usar el dropdown en el Sidebar (debajo del logo)')
doc.add_paragraph('2. O hacer clic en el selector de empresa del Topbar')
doc.add_paragraph('3. Al cambiar de empresa, se actualizan todos los datos')

doc.add_page_break()

doc.add_heading('4. GESTIÓN DE USUARIOS Y ACCESOS', 1)
doc.add_paragraph('Nota: Solo usuarios con rol Administrador pueden gestionar usuarios.')

doc.add_heading('4.1 Crear usuario (Admin)', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Ir a la página de Empresas')
doc.add_paragraph('2. Hacer clic en "+ Nuevo usuario" (visible solo para admins)')
doc.add_paragraph('3. Completar el formulario:')
doc.add_paragraph('   - Nombre de usuario *')
doc.add_paragraph('   - Email *')
doc.add_paragraph('   - Nombre completo')
doc.add_paragraph('   - Contraseña * (mínimo 6 caracteres)')
doc.add_paragraph('   - Marcar "Administrador global" si corresponde')
doc.add_paragraph('4. Hacer clic en "Crear usuario"')

doc.add_heading('4.2 Gestionar accesos por empresa', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En la tarjeta de una empresa, hacer clic en "Usuarios"')
doc.add_paragraph('2. Se expande el panel de accesos')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Agregar acceso:').bold = True
doc.add_paragraph('   a. Ingresar nombre de usuario')
doc.add_paragraph('   b. Seleccionar rol: Administrador / Editor / Visualizador')
doc.add_paragraph('   c. Hacer clic en "Agregar"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Cambiar rol:').bold = True
doc.add_paragraph('   a. Usar el dropdown junto al usuario')
doc.add_paragraph('   b. Seleccionar nuevo rol')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Remover acceso:').bold = True
doc.add_paragraph('   a. Hacer clic en "Remover" junto al usuario')

doc.add_page_break()

doc.add_heading('5. DASHBOARD - VISTA GENERAL', 1)
doc.add_paragraph('Acceso: Menú lateral → "Dashboard"')

doc.add_heading('5.1 Indicadores principales (KPIs)', 2)
doc.add_paragraph('Se muestran 4 KPIs con colores según nivel:')
doc.add_paragraph('   - VERDE (≥75%): Cumplimiento bueno')
doc.add_paragraph('   - AMARILLO (≥50%): Cumplimiento moderado')
doc.add_paragraph('   - ROJO (<50%): Cumplimiento bajo')
doc.add_paragraph()
doc.add_paragraph('   1. Total de procesos: Cantidad total de RATs')
doc.add_paragraph('   2. Completitud promedio: Porcentaje promedio de campos completados')
doc.add_paragraph('   3. Con datos sensibles: Procesos que tratan datos sensibles')
doc.add_paragraph('   4. Requieren EIPD: Procesos que necesitan Evaluación de Impacto')

doc.add_heading('5.2 Indicadores de riesgo adicionales', 2)
doc.add_paragraph('   - EIPDs pendientes')
doc.add_paragraph('   - Transferencias sin garantías')
doc.add_paragraph('   - Interés legítimo sin test')
doc.add_paragraph('   - Encargados sin contrato')

doc.add_heading('5.3 Gráfico de estados', 2)
doc.add_paragraph('Gráfico de barras verticales que muestra la distribución de RATs por estado:')
doc.add_paragraph('   - Borrador (amarillo)')
doc.add_paragraph('   - Completo (verde)')
doc.add_paragraph('   - En revisión (azul)')
doc.add_paragraph('   - Aprobado (violeta)')

doc.add_heading('5.4 Alertas de cumplimiento', 2)
doc.add_paragraph('Banners que aparecen automáticamente según condiciones:')
doc.add_paragraph('   1. Procesos con datos sensibles → Warning')
doc.add_paragraph('   2. Procesos que requieren EIPD → Peligro')
doc.add_paragraph('   3. Transferencias internacionales → Info')
doc.add_paragraph('   4. Procesos sin actualizar en +6 meses → Warning')
doc.add_paragraph('   5. Completitud promedio < 60% → Warning')
doc.add_paragraph('   6. EIPDs pendientes → Peligro')
doc.add_paragraph('   7. Transferencias sin garantías SCC/BCR → Warning')
doc.add_paragraph('   8. Interés legítimo sin test de 3 pasos → Warning')
doc.add_paragraph('   9. Encargados sin contrato Art. 14 quater → Info')

doc.add_heading('5.5 Procesos recientes', 2)
doc.add_paragraph('Tabla con los últimos 6 RATs ordenados por fecha de actualización.')

doc.add_page_break()

doc.add_heading('6. GESTIÓN DE PROCESOS RAT', 1)
doc.add_paragraph('Acceso: Menú lateral → "Procesos RAT"')

doc.add_heading('6.1 Listar y filtrar procesos', 2)
p = doc.add_paragraph()
p.add_run('Filtros disponibles:').bold = True
doc.add_paragraph('   - Por estado: Todos / Completo / Borrador / En revisión / Aprobado')
doc.add_paragraph('   - Por datos sensibles: Todos / Con datos sensibles / Sin datos sensibles')
doc.add_paragraph('   - Buscar: texto libre por nombre o categoría')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Indicadores de riesgo (columna derecha):').bold = True
doc.add_paragraph('   - ⏰ Revisar: sin actualizar +6 meses')
doc.add_paragraph('   - ⚠️ Datos sensibles')
doc.add_paragraph('   - 📋 EIPD requerida')
doc.add_paragraph('   - 🌐 Transferencia internacional')
doc.add_paragraph('   - 🤖 Decisiones automatizadas')

doc.add_heading('6.2 Crear nuevo proceso RAT (Wizard 4 pasos)', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Hacer clic en "+ Nuevo proceso" (header)')
doc.add_paragraph('2. SELECCIONAR TIPO (opcional): usar dropdown de sugerencias inteligentes')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PASO 1 - Identificación:').bold = True
doc.add_paragraph('   - Nombre del proceso *')
doc.add_paragraph('   - Categorías de titulares * (clientes, empleados, etc.)')
doc.add_paragraph('   - Fuente de los datos *')
doc.add_paragraph('   - Destinatarios/Encargados')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PASO 2 - Datos tratados:').bold = True
doc.add_paragraph('   - Categoría de datos tratados *')
doc.add_paragraph('   - Datos sensibles? (checkbox → seleccionar tipo de 7 opciones)')
doc.add_paragraph('   - Requiere EIPD? (checkbox)')
doc.add_paragraph('   - Decisiones automatizadas? (checkbox)')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PASO 3 - Finalidad y ley:').bold = True
doc.add_paragraph('   - Finalidad del tratamiento *')
doc.add_paragraph('   - Base legal * (seleccionar de 7 opciones):')
doc.add_paragraph('       • Consentimiento del titular (Art. 12)')
doc.add_paragraph('       • Ejecución de contrato (Art. 13 b)')
doc.add_paragraph('       • Obligación legal (Art. 13 a)')
doc.add_paragraph('       • Interés legítimo (Art. 16) → requiere test de 3 pasos')
doc.add_paragraph('       • Interés vital del titular (Art. 13 c)')
doc.add_paragraph('       • Datos biométricos (Art. 16 BIS)')
doc.add_paragraph('   - Test de interés legítimo (si aplica)')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PASO 4 - Almacenamiento:').bold = True
doc.add_paragraph('   - Plazo de retención *')
doc.add_paragraph('   - Medidas de seguridad')
doc.add_paragraph('   - Transferencias o comunicaciones de datos')
doc.add_paragraph('   - Transferencia internacional? (checkbox → país + garantías)')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('FINALIZAR:').bold = True
doc.add_paragraph('   3. Revisar resumen expandible')
doc.add_paragraph('   4. Hacer clic en "Guardar en el RAT"')

doc.add_heading('6.3 Editar proceso RAT', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En la tabla, hacer clic en una fila para expandir')
doc.add_paragraph('2. Hacer clic en "Editar"')
doc.add_paragraph('3. Se abre formulario con 4 pasos pre-poblados')
doc.add_paragraph('4. Modificar campos necesarios')
doc.add_paragraph('5. Cambiar estado si es necesario')
doc.add_paragraph('6. Hacer clic en "Guardar cambios"')

doc.add_heading('6.4 Duplicar proceso RAT', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Expandir la fila del RAT')
doc.add_paragraph('2. Hacer clic en "Duplicar"')
doc.add_paragraph('3. Se crea una copia con prefijo "Copia de..."')
doc.add_paragraph('4. Si requería EIPD, se marca como "pendiente"')

doc.add_heading('6.5 Eliminar proceso RAT', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Expandir la fila del RAT')
doc.add_paragraph('2. Hacer clic en "Eliminar"')
doc.add_paragraph('3. Confirmar en el diálogo')
doc.add_paragraph('4. El RAT se elimina permanentemente')

doc.add_heading('6.6 Ver historial de auditoría', 2)
doc.add_paragraph('El historial se muestra automáticamente al expandir un RAT.')
doc.add_paragraph('Contiene: acción realizada, usuario, fecha/hora.')

doc.add_page_break()

doc.add_heading('7. GESTIÓN DE BRECHAS DE SEGURIDAD', 1)
doc.add_paragraph('Acceso: Menú lateral → "Brechas"')
doc.add_paragraph('Nota: Las brechas corresponden al Art. 14 bis de la Ley 21.719.')

doc.add_heading('7.1 Ver lista de brechas', 2)
doc.add_paragraph('Tarjetas con borde según estado:')
doc.add_paragraph('   - ROJO: Plazo APDC vencido')
doc.add_paragraph('   - VERDE: Notificada APDC')
doc.add_paragraph('   - AMARILLO: Pendiente de notificación')
doc.add_paragraph()
doc.add_paragraph('Información mostrada:')
doc.add_paragraph('   - Descripción de la brecha')
doc.add_paragraph('   - Fecha de detección')
doc.add_paragraph('   - Plazo APDC (badge dinámico)')
doc.add_paragraph('   - Datos comprometidos')
doc.add_paragraph('   - RATs afectados')
doc.add_paragraph('   - Medidas adoptadas')
doc.add_paragraph('   - Estado de notificaciones (APDC y titulares)')

doc.add_heading('7.2 Registrar nueva brecha', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Hacer clic en "+ Registrar brecha" (botón rojo en header)')
doc.add_paragraph('2. Completar el formulario:')
doc.add_paragraph('   - Descripción de la brecha *')
doc.add_paragraph('   - Fecha y hora de detección * (pre-poblado con hora actual)')
doc.add_paragraph('   - Procesos RAT afectados')
doc.add_paragraph('   - Datos comprometidos')
doc.add_paragraph('   - Medidas adoptadas')
doc.add_paragraph('   - Notificado a la APDC? (checkbox)')
doc.add_paragraph('   - Notificado a los titulares? (checkbox)')
doc.add_paragraph('3. Hacer clic en "Registrar"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ALERTA: ').bold = True
p.add_run('Según la ley, la notificación a la APDC debe realizarse en 72 horas.')

doc.add_heading('7.3 Editar brecha', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En la tarjeta de la brecha, hacer clic en "Editar"')
doc.add_paragraph('2. Modificar los campos necesarios')
doc.add_paragraph('3. Actualizar estado de notificaciones si corresponde')
doc.add_paragraph('4. Hacer clic en "Guardar"')

doc.add_heading('7.4 Eliminar brecha', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En la tarjeta, hacer clic en "Eliminar"')
doc.add_paragraph('2. Confirmar en el diálogo')
doc.add_paragraph('3. La brecha se elimina permanentemente')

doc.add_page_break()

doc.add_heading('8. REPORTES Y EXPORTACIONES', 1)

doc.add_heading('8.1 Generar reportes con filtros', 2)
doc.add_paragraph('Acceso: Menú lateral → "Reportes"')
p = doc.add_paragraph()
p.add_run('Filtros disponibles:').bold = True
doc.add_paragraph('   - Buscar por nombre (texto libre)')
doc.add_paragraph('   - Estado (Todos / Borrador / Completo / En revisión / Aprobado)')
doc.add_paragraph('   - Base legal (7 opciones)')
doc.add_paragraph('   - Categoría titulares (texto libre)')
doc.add_paragraph('   - Datos sensibles (toggle)')
doc.add_paragraph('   - Requieren EIPD (toggle)')
doc.add_paragraph('   - Transferencia internacional (toggle)')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Aplicar los filtros deseados')
doc.add_paragraph('2. Hacer clic en "Aplicar filtros"')
doc.add_paragraph('3. Ver resultados en tabla')
doc.add_paragraph('4. Para limpiar: clic en "Limpiar todos"')

doc.add_heading('8.2 Exportar a CSV', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Ir a la página de Procesos RAT')
doc.add_paragraph('2. Ir a la sección inferior de la tabla')
doc.add_paragraph('3. Hacer clic en "Exportar CSV"')
doc.add_paragraph('4. Se descarga archivo: RAT_{rutEmpresa}.csv')

doc.add_heading('8.3 Exportar a PDF', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. Ir a la página de Procesos RAT')
doc.add_paragraph('2. Ir a la sección inferior de la tabla')
doc.add_paragraph('3. Hacer clic en "Exportar PDF"')
doc.add_paragraph('4. Se descarga archivo: RAT_{rutEmpresa}.pdf')
doc.add_paragraph('Nota: El PDF incluye todos los campos legales del RAT.')

doc.add_page_break()

doc.add_heading('9. CONFIGURACIÓN DE CUENTA', 1)

doc.add_heading('9.1 Cambiar contraseña', 2)
p = doc.add_paragraph()
p.add_run('Pasos:').bold = True
doc.add_paragraph('1. En el Sidebar, hacer clic en "Cambiar contraseña"')
doc.add_paragraph('2. Se abre modal con formulario:')
doc.add_paragraph('   - Contraseña actual *')
doc.add_paragraph('   - Nueva contraseña * (mínimo 6 caracteres)')
doc.add_paragraph('   - Confirmar nueva contraseña *')
doc.add_paragraph('3. Hacer clic en "Cambiar"')
doc.add_paragraph('4. Se muestra toast de éxito o error')

doc.add_heading('9.2 Navegación rápida', 2)
p = doc.add_paragraph()
p.add_run('Sidebar:').bold = True
doc.add_paragraph('   - Logo Custodio + "Ley 21.719"')
doc.add_paragraph('   - Selector de empresa activa')
doc.add_paragraph('   - Dashboard')
doc.add_paragraph('   - Procesos RAT')
doc.add_paragraph('   - Reportes')
doc.add_paragraph('   - Brechas')
doc.add_paragraph('   - Empresas')
doc.add_paragraph('   - Cambiar contraseña')
doc.add_paragraph('   - Cerrar sesión')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Topbar:').bold = True
doc.add_paragraph('   - Selector de empresa activa (con buscador)')
doc.add_paragraph('   - Badge "Admin" (si corresponde)')
doc.add_paragraph('   - Nombre del usuario')

doc.add_page_break()

doc.add_heading('ANEXO: ROLES Y PERMISOS', 1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Rol'
hdr_cells[1].text = 'Alcance'
hdr_cells[2].text = 'Permisos'

hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[1].paragraphs[0].runs[0].bold = True
hdr_cells[2].paragraphs[0].runs[0].bold = True

row1 = table.rows[1].cells
row1[0].text = 'Administrador'
row1[1].text = 'Global'
row1[2].text = ' CRUD empresas, usuarios, RATs, brechas'

row2 = table.rows[2].cells
row2[0].text = 'Editor'
row2[1].text = 'Por empresa'
row2[2].text = ' CRUD RATs y brechas de su empresa'

row3 = table.rows[3].cells
row3[0].text = 'Visualizador'
row3[1].text = 'Por empresa'
row3[2].text = ' Solo lectura de RATs y brechas'

doc.add_paragraph()
doc.add_heading('ANEXO: CAMPOS OBLIGATORIOS DEL RAT (Art. 16)', 1)

table2 = doc.add_table(rows=10, cols=2)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table2.rows[0].cells
hdr[0].text = 'Campo'
hdr[1].text = 'Obligatorio'
hdr[0].paragraphs[0].runs[0].bold = True
hdr[1].paragraphs[0].runs[0].bold = True

fields = [
    ('Nombre del proceso', 'Sí'),
    ('Categoría de datos', 'Sí'),
    ('Categorías de titulares', 'Sí'),
    ('Finalidad', 'Sí'),
    ('Base legal', 'Sí'),
    ('Fuente de datos', 'Sí'),
    ('Plazo de retención', 'Sí'),
    ('Destinatarios/Encargados', 'No'),
    ('Medidas de seguridad', 'No'),
]

for i, (field, req) in enumerate(fields, 1):
    row = table2.rows[i].cells
    row[0].text = field
    row[1].text = req

doc.add_paragraph()
doc.add_paragraph('Versión del documento: 1.0')
doc.add_paragraph('Sistema: Custodio RAT Manager · Ley 21.719')

output_path = 'C:/Users/chelo/Desktop/RAT_opencode/docs/Manual_Usuario_Custodio.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')