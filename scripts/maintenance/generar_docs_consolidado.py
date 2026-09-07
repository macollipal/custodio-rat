"""
generar_docs_consolidado.py
============================
Genera documentos Word v1.13 COMPLETAMENTE ACUMULATIVOS desde v1.0.

Estrategia:
  1. Ejecuta los scripts v1.10 (base acumulativa más completa disponible)
     en su propio directorio para que importen _theme_custodio correctamente.
  2. Abre cada .docx resultante con python-docx.
  3. Inserta secciones con el contenido incremental de v1.11 y v1.12.
  4. Guarda como v1.13 en docs/documentacion_oficial/.
  5. Elimina los archivos intermedios v1.10 para no dejar duplicados.

Documentos generados:
  02_Requisitos, 03_Historias_Usuario, 09_Backlog_Producto, 10_Plan_QA
  (04, 06, 08, 12, MTX se actualizan por separado cuando se disponga
   de sus scripts acumulativos v1.10)
"""

import sys
import os
import subprocess
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Rutas ─────────────────────────────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent.parent.parent   # RAT_opencode/
OUT_DIR = ROOT / "docs" / "documentacion_oficial"
AUD_V110 = ROOT / "docs" / "auditorias" / "2026-07-18_auditoria_doc_v1.10" / "_scripts"
AUD_V17  = ROOT / "docs" / "auditorias" / "2026-06-24_auditoria_v1.7" / "_scripts"
AUD_V111 = ROOT / "docs" / "auditorias" / "2026-08-22_auditoria_qa_tests" / "_scripts"
AUD_V19  = ROOT / "docs" / "auditorias" / "2026-07-05_auditoria_v1.9" / "_scripts"

VERSION = "v1.13"
COLOR_PRI = RGBColor(0x1F, 0x49, 0x7D)   # azul Custodio

# ── Utilidades python-docx ────────────────────────────────────────────────────

def _shading(cell, fill_hex="1F497D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths_cm=None, shade_header=True):
    """Tabla con cabecera sombreada."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # cabecera
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        if shade_header:
            _shading(cell)
    # filas
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r + 1].cells[c]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    # anchos opcionales
    if col_widths_cm:
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                if c_idx < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[c_idx])
    return t


def add_italic_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def section_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

# ── Runner de script v1.10 ────────────────────────────────────────────────────

def run_script(script_path: Path) -> bool:
    """Ejecuta un script build en su propio directorio."""
    print(f"  → {script_path.name}...")
    res = subprocess.run(
        [sys.executable, str(script_path)],
        cwd=str(script_path.parent),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if res.returncode != 0:
        print(f"    ERROR: {res.stderr[-400:]}")
        return False
    print(f"    OK: {(res.stdout or '').strip()[-120:]}")
    return True


def find_docx(prefix: str):
    hits = sorted(OUT_DIR.glob(f"{prefix}*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    return hits[0] if hits else None


def rename_to_v113(path: Path) -> Path:
    """Renombra un .docx quitando su versión e insertando v1.13."""
    stem = path.stem
    # Quita _v1.xx al final si existe
    import re
    stem_clean = re.sub(r"_v\d+\.\d+$", "", stem)
    new_name = OUT_DIR / f"{stem_clean}_{VERSION}.docx"
    return new_name


# ══════════════════════════════════════════════════════════════════════════════
# 02 — Requisitos
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_02_requisitos():
    print("\n[02] Requisitos...")

    # 1. Base acumulativa v1.10 (RF-001 a RF-169 + todas las RNF)
    ok = run_script(AUD_V110 / "build_02_requisitos_v1_10.py")
    if not ok:
        return

    base = find_docx("02_Requisitos")
    if not base:
        print("  WARN: no se encontró 02_Requisitos*.docx")
        return

    doc = Document(str(base))

    # 2. Nuevos en v1.12 — RF-174 a RF-179
    section_heading(doc, "Requisitos Funcionales v1.12 (Septiembre 2026)", level=1)
    doc.add_paragraph(
        "Los siguientes requisitos funcionales fueron añadidos en la iteración v1.12:"
    )
    rf_v112 = [
        ["RF-174", "Alta", "ARCO público", "GET /publico/verificar-titular: banner amarillo si el email ya tiene un TKT abierto (Art. 12)."],
        ["RF-175", "Alta", "ARCO público", "POST /tkt-solicitud-derecho/: acuse de recibo automático al titular con código de seguimiento y plazo (email en ≤5 min)."],
        ["RF-176", "Media", "ARCO interno", "TicketDrawer.tsx: chips de placeholders dinámicos bajo el textarea de respuesta para agilizar redacción."],
        ["RF-177", "Alta", "ARCO interno", "FlujoModal.tsx: semáforo SLA visual (verde/amarillo/rojo) con días hábiles transcurridos y restantes."],
        ["RF-178", "Alta", "Empresas", "CompanyFichaPanel.tsx: ficha completa de empresa con 4 tabs lazy (RATs, ARCO, Brechas, Stats) — Art. 16 Ley 21.719."],
        ["RF-179", "Alta", "DevSecOps", "CI/CD pip-audit: cero vulnerabilidades críticas en dependencias Python antes de cada deploy."],
    ]
    add_table(doc, ["RF", "Prioridad", "Módulo", "Descripción"], rf_v112,
              col_widths_cm=[1.5, 1.5, 2.5, 12.0])

    # 3. Nuevos en v1.12 — RNF-21, RNF-22
    section_heading(doc, "Requisitos No Funcionales v1.12", level=2)
    rnf_v112 = [
        ["RNF-21", "Global (código, tests, docs)", "Nomenclatura APDP: usar 'titular' en lugar de 'solicitante' en toda la aplicación (art. 12 Ley 21.719)."],
        ["RNF-22", "GET /publico/verificar-titular", "Rate-limit: máximo 10 peticiones por hora por IP para prevenir abuso de la búsqueda por email."],
    ]
    add_table(doc, ["RNF", "Alcance", "Descripción"], rnf_v112,
              col_widths_cm=[1.5, 4.0, 12.0])

    # 4. Guardar como v1.13
    dest = rename_to_v113(base)
    # Eliminar v1.13 anterior si existe
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    # Limpiar intermedio v1.10 si es diferente al destino
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# 03 — Historias de Usuario
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_03_historias():
    print("\n[03] Historias de Usuario...")

    # 1. Base acumulativa: v1.7 tiene HU-001 a HU-085 (mejor base cumplativa)
    ok = run_script(AUD_V17 / "build_03_historias_usuario_v1_7.py")
    if not ok:
        return

    base = find_docx("03_Historias")
    if not base:
        print("  WARN: no se encontró 03_Historias*.docx")
        return

    doc = Document(str(base))

    # 2. HU-064 a HU-071 (v1.5 Seguridad + v1.6 UI/UX — ausentes en tabla v1.7)
    section_heading(doc, "HU-064 a HU-071 — Seguridad y UI/UX (v1.5 y v1.6)", level=1)
    add_italic_note(doc,
        "Estas HUs corresponden a las iteraciones v1.5 (Seguridad) y v1.6 (UI/UX). "
        "Completan el listado entre HU-063 y HU-072."
    )
    hus_v15_v16 = [
        ["HU-064", "EP-01", "RNF-114", "Alta", "M",  "Implementar CSRF middleware en todos los endpoints públicos (Art. seguridad)"],
        ["HU-065", "EP-01", "RNF-115", "Alta", "L",  "Encryption at Rest: cifrado Fernet para datos sensibles en BD (BYTEA)"],
        ["HU-066", "EP-02", "RNF-116", "Alta", "M",  "Service Layer: extraer lógica de negocio de routes a services/"],
        ["HU-067", "EP-02", "RNF-040", "Media", "S", "Schemas Pydantic v2: separar Request/Response en schemas independientes"],
        ["HU-068", "EP-02", "RF-029",  "Alta", "M",  "RatDetailModal: modal de detalle RAT con vista completa de 25 campos"],
        ["HU-069", "EP-05", "RF-064",  "Alta", "S",  "Drawer ARCO responsive: diseño adaptado a móvil con breakpoints Tailwind"],
        ["HU-070", "EP-06", "RF-074",  "Media", "S", "Dashboard clickable: cards del dashboard navegan a filtros pre-aplicados"],
        ["HU-071", "EP-02", "RF-027",  "Baja", "S",  "Sort estable en tablas: mantiene orden secundario al reordenar columnas"],
    ]
    add_table(doc, ["HU", "Épica", "Trazabilidad RF/RNF", "Prioridad", "Tamaño", "Título"],
              hus_v15_v16, col_widths_cm=[1.3, 1.3, 2.0, 1.4, 1.2, 10.3])

    # 3. HU-086 a HU-103 (v1.10 — Iter 11+12 y fixes v1.9)
    section_heading(doc, "HU-086 a HU-103 — Iter 11+12 y Fixes (v1.10)", level=1)
    hus_v110 = [
        # Iter 11: 15 campos Tier 1+2
        ["HU-086", "EP-02", "RF-141",      "Alta",   "M", "Registrar datos NNA en RAT (datos_nna, datos_anonimizados, datos_seudonimizados)"],
        ["HU-087", "EP-02", "RF-142–149",  "Media",  "S", "Registrar nivel de confidencialidad y estructura del dato"],
        ["HU-088", "EP-02", "RF-146–149",  "Media",  "S", "Registrar ciclo de procesamiento y frecuencia de tratamiento"],
        ["HU-089", "EP-02", "RF-150–155",  "Media",  "M", "Registrar campos operativos Tier 2 (sistema_almacenamiento, volumen, email, etc.)"],
        # Iter 12: fixes críticos
        ["HU-090", "EP-02", "RF-156",      "CRÍTICO","S", "Validar límite 10 MB en archivos BYTEA (archivo_base_legal, tkt_adjunto)"],
        ["HU-091", "EP-02", "RF-157",      "CRÍTICO","S", "Test de interés legítimo con mínimo 50 caracteres (Art. 16 Ley 21.719)"],
        ["HU-092", "EP-05", "RF-158",      "CRÍTICO","M", "Hash SHA-256 de evidencia ARCO computado automáticamente al resolver TKT"],
        ["HU-093", "EP-05", "RF-159",      "ALTO",   "S", "causal_rechazo con enum cerrado de 7 causales Art. 29 RL"],
        ["HU-094", "EP-05", "RF-160",      "ALTO",   "S", "Toggle ARCO con touch target 44×44px para accesibilidad móvil (WCAG 2.1)"],
        ["HU-095", "EP-04", "RF-161",      "ALTO",   "M", "Notificación APDC automatizada al marcar notificado_apdc=true en brecha"],
        ["HU-096", "EP-04", "RF-162",      "ALTO",   "M", "Notificación a titulares automatizada al marcar notificado_titulares=true"],
        ["HU-097", "EP-05", "RF-158",      "ALTO",   "S", "TKT no puede resolverse sin evidencia ni hash SHA-256 (HTTP 400)"],
        # v1.10 fixes (Julio 2026)
        ["HU-098", "EP-02", "RF-163",      "CRÍTICO","M", "IDOR multi-tenant: get_rat_for_user() retorna 404 si empresa no coincide; superadmin accede a todos"],
        ["HU-099", "EP-02", "RF-164",      "ALTO",   "S", "base_legal_valida strict: validación contra enum taxativo de 6 opciones"],
        ["HU-100", "EP-11", "RF-165",      "ALTO",   "S", "ConsentimientoAlert: listarConsentimientos() antes de guardar si datos_sensibles=True"],
        ["HU-101", "EP-02", "RF-166",      "ALTO",   "L", "Homologación orden campos RAT entre RatDetailView, RatEditForm, RatWizard y PDF (5 pasos canónicos)"],
        ["HU-102", "EP-06", "RF-167",      "MEDIO",  "M", "PDF con títulos de sección (PASO 1 — IDENTIFICACIÓN, etc.) con fondo COLOR_PRIMARIO"],
        ["HU-103", "EP-02", "RF-168",      "MEDIO",  "S", "Encoding UTF-8 corregido en backend: ALERTAS_AUDITORIA, regex, conftest.py, main.py"],
    ]
    add_table(doc, ["HU", "Épica", "Trazabilidad", "Prioridad", "Tamaño", "Título"],
              hus_v110, col_widths_cm=[1.3, 1.3, 2.0, 1.4, 1.2, 10.3])

    # 4. HU-104 a HU-108 (v1.12)
    section_heading(doc, "HU-104 a HU-108 — ARCO y Empresa (v1.12)", level=1)
    hus_v112 = [
        ["HU-104", "EP-05", "RF-174", "Alta",  "S", "Titular: detección de solicitud ARCO duplicada antes de enviar (banner amarillo)"],
        ["HU-105", "EP-05", "RF-175", "Alta",  "S", "Titular: acuse de recibo inmediato con código de seguimiento y plazo (email ≤5 min)"],
        ["HU-106", "EP-05", "RF-176", "Media", "S", "Operador DPO: insertar placeholders dinámicos en respuesta al titular con un clic"],
        ["HU-107", "EP-05", "RF-177", "Alta",  "S", "Operador DPO: semáforo SLA visual verde/amarillo/rojo con días hábiles transcurridos"],
        ["HU-108", "EP-01", "RF-178", "Alta",  "M", "Admin empresa: ficha completa con 4 tabs lazy (RATs, ARCO, Brechas, Stats)"],
    ]
    add_table(doc, ["HU", "Épica", "Trazabilidad", "Prioridad", "Tamaño", "Título"],
              hus_v112, col_widths_cm=[1.3, 1.3, 2.0, 1.4, 1.2, 10.3])

    # 5. Guardar como v1.13
    dest = rename_to_v113(base)
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# 09 — Backlog de Producto
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_09_backlog():
    print("\n[09] Backlog de Producto...")

    # 1. Base acumulativa v1.10 (items v1.0 a v1.9)
    ok = run_script(AUD_V110 / "build_09_backlog_v1_10.py")
    if not ok:
        return

    base = find_docx("09_Backlog")
    if not base:
        print("  WARN: no se encontró 09_Backlog*.docx")
        return

    doc = Document(str(base))

    # 2. Items nuevos v1.12 (cerrados + pendientes actualizados)
    section_heading(doc, "Items Cerrados en v1.12 (Septiembre 2026)", level=1)
    items_v112_cerrados = [
        ["TKT-V112-01", "P1", "Feature", "Detección duplicado ARCO: GET /publico/verificar-titular — RF-174", "CERRADO"],
        ["TKT-V112-02", "P1", "Feature", "Acuse de recibo automático al titular: email en ≤5 min — RF-175", "CERRADO"],
        ["TKT-V112-03", "P2", "Feature", "Chips placeholders en TicketDrawer.tsx — RF-176", "CERRADO"],
        ["TKT-V112-04", "P1", "Feature", "Semáforo SLA en FlujoModal.tsx — RF-177", "CERRADO"],
        ["TKT-V112-05", "P1", "Feature", "Ficha empresa con 4 tabs lazy (CompanyFichaPanel.tsx) — RF-178", "CERRADO"],
        ["TKT-V112-06", "P1", "DevSecOps", "pip-audit en CI/CD: 0 vulns críticas Python — RF-179", "CERRADO"],
        ["TKT-V112-07", "P2", "Nomenclatura", "Renombrar 'solicitante' → 'titular' en toda la app — RNF-21", "CERRADO"],
        ["TKT-V112-08", "P2", "Seguridad", "Rate-limit en GET /publico/verificar-titular — RNF-22", "CERRADO"],
    ]
    add_table(doc, ["ID", "Prioridad", "Tipo", "Título", "Estado"],
              items_v112_cerrados, col_widths_cm=[2.2, 1.5, 2.0, 10.5, 1.7])

    section_heading(doc, "Items Pendientes Actualizados (v1.12)", level=1)
    items_pendientes = [
        ["QW-ITER14-01", "P2", "Feature", "Paginación en listados >100 registros (RAT/ARCO/Brechas/Encargados)"],
        ["QW-ITER14-02", "P3", "Feature", "Retry logic en OCI uploads (resiliencia ante timeouts)"],
        ["QW-ITER14-03", "P2", "Feature", "Logs de auditoría en tabla audit_log independiente (Art. 28 Ley 21.719)"],
        ["QW-ITER14-04", "P3", "Feature", "ALTER TABLE categoria_titulares SET NOT NULL (requiere migración coordinada)"],
    ]
    add_table(doc, ["ID", "Prioridad", "Tipo", "Título"],
              items_pendientes, col_widths_cm=[2.2, 1.5, 2.0, 12.2])

    dest = rename_to_v113(base)
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# 10 — Plan de QA
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_10_plan_qa():
    print("\n[10] Plan de QA...")

    # 1. Base v1.10 (TC-030 a TC-046 — previos en documento v1.9 ya embebidos)
    ok = run_script(AUD_V110 / "build_10_plan_qa_v1_10.py")
    if not ok:
        return

    base = find_docx("10_Plan_QA")
    if not base:
        print("  WARN: no se encontró 10_Plan_QA*.docx")
        return

    doc = Document(str(base))

    # 2. TC-047 a TC-055 (v1.11 — QA Total 2026-08-22)
    section_heading(doc, "Casos de Prueba v1.11 — QA Total (22 Ago 2026)", level=1)
    add_italic_note(doc, "v1.11: QA Total completado. 78 fallos → 0. Suite final: 732 tests, 0 fallos.")
    tc_v111 = [
        ["TC-047", "CRÍTICO", "Backend",  "PATCH TKT resuelto CON metodo_verificacion_identidad → 200",
         "PATCH /tkt-solicitud-derecho/{id} con estado=resuelto + metodo_verificacion_identidad",
         "HTTP 200, estado=resuelto", "pytest"],
        ["TC-048", "CRÍTICO", "Backend",  "Prorrogar TKT desde estado resuelto → 400",
         "POST /tkt-solicitud-derecho/{id}/prorrogar (ticket resuelto)",
         "HTTP 400", "pytest"],
        ["TC-049", "ALTO",    "Backend",  "POST /auth/users retorna 201 al crear usuario nuevo",
         "POST /auth/users con datos válidos",
         "HTTP 201 con objeto usuario", "pytest"],
        ["TC-050", "ALTO",    "Backend",  "encrypt_existing_bytea falla si ENCRYPTION_KEY='' (sin fallback)",
         "Ejecutar _check_prerequisites() con ENCRYPTION_KEY=''",
         "SystemExit", "pytest"],
        ["TC-051", "ALTO",    "Backend",  "encrypt_existing_bytea detecta datos ya cifrados como Fernet",
         "is_already_encrypted(fernet.encrypt(b'test'))",
         "True", "pytest"],
        ["TC-052", "ALTO",    "Backend",  "encrypt_existing_bytea dry-run no modifica BD",
         "_migrate_table(..., dry_run=True)",
         "stats['migrados']==0, datos sin cambios", "pytest"],
        ["TC-053", "ALTO",    "Backend",  "encrypt_existing_bytea segunda pasada idempotente",
         "_migrate_table() dos veces sobre mismos datos",
         "Segunda: ya_cifrados=1, migrados=0", "pytest"],
        ["TC-054", "ALTO",    "Backend",  "EIPD validator bloquea datos_sensibles=True sin campos EIPD",
         "POST /rats/ con datos_sensibles=True sin evaluacion_impacto+estado_eipd",
         "HTTP 422", "pytest"],
        ["TC-055", "MEDIO",   "Backend",  "Ruta /auditoria/verify-chain no capturada por /{company_id}",
         "GET /rats/auditoria/verify-chain (superadmin)",
         "HTTP 200 con chain info", "pytest"],
    ]
    add_table(doc, ["ID", "Sev.", "Nivel", "Descripción", "Pasos", "Resultado Esperado", "Framework"],
              tc_v111, col_widths_cm=[1.2, 1.3, 1.5, 3.5, 4.5, 3.5, 1.5])

    # 3. TC-056 a TC-063 (v1.12)
    section_heading(doc, "Casos de Prueba v1.12 (Septiembre 2026)", level=1)
    tc_v112 = [
        ["TC-056", "ALTO",    "Backend",  "GET /publico/verificar-titular — sin TKT abierto → 200 {}",
         "GET /publico/verificar-titular?email=sin_ticket@test.com",
         "HTTP 200, tiene_ticket=false", "pytest"],
        ["TC-057", "CRÍTICO", "Backend",  "GET /publico/verificar-titular — con TKT abierto → 200 con datos",
         "GET /publico/verificar-titular?email=con_ticket@test.com",
         "HTTP 200, tiene_ticket=true + datos TKT", "pytest"],
        ["TC-058", "ALTO",    "Backend",  "GET /publico/verificar-titular — email no registrado → 200 {}",
         "GET /publico/verificar-titular?email=desconocido@test.com",
         "HTTP 200, tiene_ticket=false", "pytest"],
        ["TC-059", "ALTO",    "Backend",  "Rate-limit verificar-titular: >10/h por IP → 429",
         "11 peticiones consecutivas desde misma IP",
         "HTTP 429 en la 11.ª", "pytest + rate-limit test"],
        ["TC-060", "ALTO",    "Backend",  "POST /tkt-solicitud-derecho/ → acuse de recibo enviado (mock SMTP)",
         "POST con email válido, SMTP mockeado",
         "acuse_enviado_at != null en response", "pytest + mock"],
        ["TC-061", "MEDIO",   "Backend",  "POST TKT sin email → acuse NO enviado (sin fallo)",
         "POST con titular_email=None",
         "HTTP 201, acuse_enviado_at=null", "pytest"],
        ["TC-062", "ALTO",    "Frontend", "CompanyFichaPanel tabs cargan lazy sin error (4 tabs)",
         "Navegar a /empresas/{id}, clicar cada tab",
         "Contenido visible, sin error 500", "Playwright"],
        ["TC-063", "MEDIO",   "Global",   "Búsqueda de 'solicitante' en codebase → 0 ocurrencias (RNF-21)",
         "grep -r 'solicitante' src/ backend/",
         "0 matches", "CI grep check"],
    ]
    add_table(doc, ["ID", "Sev.", "Nivel", "Descripción", "Pasos", "Resultado Esperado", "Framework"],
              tc_v112, col_widths_cm=[1.2, 1.3, 1.5, 3.5, 4.5, 3.5, 1.5])

    # 4. Resumen de cobertura actualizado
    section_heading(doc, "Resumen de Cobertura Acumulado v1.13", level=1)
    cobertura = [
        ["Backend pytest",    "~740 tests", "pytest + httpx", "TC-001 a TC-063 cubiertos", "0 fallos"],
        ["Frontend TypeScript","0 errores",  "tsc --noEmit",  "Sin regresiones de tipos",  "0 errores"],
        ["E2E Playwright",    "~65 tests",  "Playwright",    "Flows ARCO, EIPD, Brechas",  "OK"],
        ["Integration",       "Neon QA",    "PostgreSQL",    "custodio_test DB",            "OK"],
    ]
    add_table(doc, ["Tipo", "Cantidad", "Framework", "Alcance", "Estado"],
              cobertura, col_widths_cm=[2.5, 1.5, 2.5, 5.5, 2.0])

    dest = rename_to_v113(base)
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# 04 — Casos de Uso
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_04_casos_uso():
    print("\n[04] Casos de Uso...")

    # Base: v1.7 — tiene "Listado consolidado" CU-001 a CU-068
    ok = run_script(AUD_V17 / "build_04_casos_uso_v1_7.py")
    if not ok:
        return

    base = find_docx("04_Casos_de_Uso")
    if not base:
        print("  WARN: no se encontró 04_Casos_de_Uso*.docx")
        return

    doc = Document(str(base))

    # v1.8/v1.9: CU-069 a CU-082
    section_heading(doc, "Casos de Uso v1.8–v1.9 (Iter 11+12 — Julio 2026)", level=1)
    cu_v19 = [
        ["CU-069", "Registrar 15 campos Tier 1+Tier 2 en RAT", "AC-02/03", "RF-141–155", "PUT /rats/{id}"],
        ["CU-070", "BYTEA limitado a 10MB en archivo y adjunto ARCO", "Sistema", "RF-156", "CHECK constraint PostgreSQL"],
        ["CU-071", "Test IL validado con mínimo 50 caracteres", "AC-02/03", "RF-157", "PUT /rats/{id} — validación Pydantic + frontend"],
        ["CU-072", "Hash SHA-256 automático de evidencia ARCO al resolver TKT", "Sistema", "RF-158", "PATCH /tkt-solicitud-derecho/{id}"],
        ["CU-073", "causal_rechazo con enum cerrado (7 causales Art. 29 RL)", "AC-01/02", "RF-159", "PATCH /tkt-solicitud-derecho/{id} — dropdown"],
        ["CU-074", "Toggle ARCO con touch target 44x44px (mobile)", "AC-05", "RF-160", "UI: solicitud_derecho/page.tsx"],
        ["CU-075", "Notificación APDC automatizada al crear brecha", "Sistema", "RF-161", "actualizar_brecha() — email_service"],
        ["CU-076", "Notificación a titulares automatizada al cerrar brecha", "Sistema", "RF-162", "actualizar_brecha() — logging"],
        ["CU-077", "TKT no puede resolverse sin evidencia ni hash", "AC-01/02", "RF-158", "PATCH /tkt → HTTP 400 si sin adjuntos ni respuesta"],
        ["CU-078", "IDOR multi-tenant: empresa no puede acceder a RAT de otra", "AC-02/03", "RF-163", "get_rat_for_user() — retorna 404 en 6 endpoints"],
        ["CU-079", "base_legal_valida strict contra enum taxativo", "Sistema", "RF-164", "base_legal_valida() — 6 opciones válidas"],
        ["CU-080", "ConsentimientoAlert antes de guardar RAT con datos_sensibles", "Sistema", "RF-165", "handleSave() — listarConsentimientos()"],
        ["CU-081", "Homologación orden campos RAT en wizard, drawer y PDF", "Sistema", "RF-166", "RatDetailView + RatEditForm + RatWizard + export_service"],
        ["CU-082", "PDF con títulos de sección y alertas rojas", "Sistema", "RF-167", "export_service — PASOx con COLOR_PRIMARIO"],
    ]
    add_table(doc, ["ID", "Nombre", "Actores", "Trazabilidad RF", "Disparador"], cu_v19,
              col_widths_cm=[1.5, 4.2, 2.0, 2.5, 4.8])

    # v1.12: CU-083 a CU-087 (nombrados CU-31/35 en el script original)
    section_heading(doc, "Casos de Uso v1.12 (Septiembre 2026)", level=1)
    cu_v112 = [
        ["CU-083", "Detección de titular repetido en formulario público", "AC-05 (Anónimo)", "RF-174", "GET /publico/verificar-titular — banner si ya tiene TKT abierto"],
        ["CU-084", "Acuse de recibo automático al crear ticket ARCO", "AC-02/03 (editor+)", "RF-175", "POST /tkt-solicitud-derecho/ — email en ≤5 min con tracking token"],
        ["CU-085", "Insertar placeholders dinámicos en respuesta ARCO", "AC-02/03 (editor+)", "RF-176", "Frontend TicketDrawer.tsx — chips {{nombre_titular}}, {{fecha}}, etc."],
        ["CU-086", "Ver flujo ARCO con semáforo SLA y tiempos reales", "AC-02/03 (editor+)", "RF-177", "Frontend FlujoModal.tsx — días hábiles consumidos/restantes"],
        ["CU-087", "Ficha de empresa con 4 tabs lazy (Datos/RATs/ARCO/Brechas)", "AC-01/02/03", "RF-178", "Frontend CompanyFichaPanel.tsx — Art. 16 Ley 21.719"],
    ]
    add_table(doc, ["ID", "Nombre", "Actores", "Trazabilidad RF", "Disparador"], cu_v112,
              col_widths_cm=[1.5, 4.2, 2.5, 2.0, 4.8])

    dest = rename_to_v113(base)
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# 08 — API REST
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_08_api():
    print("\n[08] API REST...")

    # Base: v1.10 — tiene endpoints PATCH TKT + IDOR multi-tenant
    ok = run_script(AUD_V110 / "build_08_api_v1_10.py")
    if not ok:
        return

    base = find_docx("08_API_REST")
    if not base:
        base = find_docx("08_API")
    if not base:
        print("  WARN: no se encontró 08_API*.docx")
        return

    doc = Document(str(base))

    # v1.11: POST /auth/users→201, GET /publico/csrf-token, PUT /transparencia, PATCH TKT metodo
    section_heading(doc, "Endpoints nuevos/modificados v1.11 (Agosto 2026)", level=1)
    ep_v111 = [
        ["POST", "/auth/users", "JWT", "superadmin",
         "UserCreate: username, email, full_name, password, rol_global",
         "201 + UserOut (antes: 200)"],
        ["GET", "/publico/csrf-token", "No", "Pública (rate 30/min)",
         "—",
         "200 + {token, header_name, expires_in_seconds}"],
        ["PUT", "/transparencia/{company_id}", "JWT", "admin_empresa, superadmin",
         "overrides_json: dict personalizado de items",
         "200 + PoliticaTransparenciaResponse"],
        ["PATCH", "/tkt-solicitud-derecho/{id}", "JWT", "editor+",
         "estado, respuesta_texto, metodo_verificacion_identidad*, causal_rechazo*",
         "200 + TktTicketResponse — metodo_verificacion ahora puede ir en body del PATCH"],
    ]
    add_table(doc, ["Método", "Path", "Auth", "RBAC", "Params", "Response"], ep_v111,
              col_widths_cm=[1.4, 4.5, 1.0, 2.5, 4.5, 4.1])

    # v1.12: GET /publico/verificar-titular + acuse recibo en POST TKT
    section_heading(doc, "Endpoints nuevos/modificados v1.12 (Septiembre 2026)", level=1)
    ep_v112 = [
        ["GET", "/publico/verificar-titular", "No", "Pública (rate 20/min)",
         "company_id, email (query params)",
         "200 + {tiene_tickets_abiertos: bool, cantidad: int}"],
        ["POST", "/publico/ejercer-derechos", "No", "Pública (10/hora/IP)",
         "EjercerDerechosRequest (company_id, tipo, nombre, email, rut...)",
         "201 + {tracking_token, mensaje} — envía acuse recibo automático (QW6)"],
        ["POST", "/tkt-solicitud-derecho/", "JWT", "editor+",
         "TktCreate",
         "201 + TktOut — envía acuse de recibo al titular_email si presente (ARCO-QW6)"],
    ]
    add_table(doc, ["Método", "Path", "Auth", "RBAC", "Params", "Response"], ep_v112,
              col_widths_cm=[1.4, 4.5, 1.0, 2.5, 4.5, 4.1])
    add_italic_note(doc, "Nota v1.12: RNF-22 — GET /publico/verificar-titular tiene rate limit 20/min por IP "
                        "para prevenir enumeración de emails. Nomenclatura APDP corregida en todo el sistema (RNF-21).")

    dest = rename_to_v113(base)
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# 12 — Manual Técnico
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_12_manual_tecnico():
    print("\n[12] Manual Técnico...")

    # Base: v1.10 — BYTEA 10MB, Test IL, Hash SHA-256, causal_rechazo, toggle 44px
    ok = run_script(AUD_V110 / "build_12_manual_tecnico_v1_10.py")
    if not ok:
        return

    base = find_docx("12_Manual_Tecnico")
    if not base:
        print("  WARN: no se encontró 12_Manual_Tecnico*.docx")
        return

    doc = Document(str(base))

    # v1.11: fixes QA total + Sprint A+B+UX
    section_heading(doc, "Cambios técnicos v1.11 (Agosto 2026 — QA Total)", level=1)
    cambios_v111 = [
        ["POST /auth/users → 201", "ALTO", "routes/auth.py:129",
         "Agregado status_code=201. Antes retornaba 200 por default FastAPI."],
        ["PATCH TKT: metodo_verificacion en body", "ALTO", "routes/tkt_solicitud_derecho.py",
         "Condición corregida: verifica campo BD Y body del PATCH. Antes ignoraba el body."],
        ["encrypt_existing_bytea: ENCRYPTION_KEY estricta", "MEDIO", "scripts/migration/encrypt_existing_bytea.py",
         "Eliminado fallback a settings.ENCRYPTION_KEY. Exige key explícita en entorno."],
        ["Sprint A: soft delete RAT (Art. 19)", "ALTO", "services/rat_crud.py, routes/rats.py",
         "delete_rat() asigna deleted_at. Filtros excluyen deleted_at IS NOT NULL. RATs aprobados no eliminables (409)."],
        ["Sprint A: EIPD gate aprobación (Art. 15 bis)", "ALTO", "services/rat_crud.py (aprobar_rat)",
         "validar_eipd_obligatoria() bloquea APROBADO sin EIPD cuando datos_sensibles=True o transferencia_int=True."],
        ["Sprint A: datos NNA (Art. 16 §4)", "ALTO", "models/rat.py, schemas/rat.py",
         "Campo datos_nna boolean. Bloqueo aprobación si datos_nna=True sin EIPD."],
        ["Sprint A: mutex anon/seudo", "MEDIO", "schemas/rat.py",
         "No se puede marcar datos_anonimizados=True y datos_seudonimizados=True simultáneamente (validator Pydantic)."],
        ["Sprint B: respuesta_texto obligatoria (Art. 12)", "ALTO", "services/ticket_service.py",
         "Bloquea PATCH→resuelto sin respuesta_texto no vacía."],
        ["Sprint B: SLA alerts 72h brechas", "ALTO", "services/breach_service.py",
         "Alerta automática APDC si brecha sin notificación tras 72h (Art. 14 bis §3)."],
        ["Sprint UX: política transparencia editable (Art. 14 ter)", "MEDIO", "routes/transparencia.py",
         "PUT /transparencia/{company_id} con overrides_json. Hash SHA-256 recalculado en cada guardado."],
    ]
    add_table(doc, ["Cambio", "Severidad", "Archivos", "Descripción"], cambios_v111,
              col_widths_cm=[3.0, 1.5, 4.0, 9.5])

    # v1.12: QW5, QW6, QW7, QW8, Empresas-QW6, CI/CD, Nomenclatura
    section_heading(doc, "Cambios técnicos v1.12 (Septiembre 2026)", level=1)
    cambios_v112 = [
        ["Público-QW5: verificar-titular", "MEDIO", "routes/publico_arco.py",
         "GET /publico/verificar-titular. Rate 20/min. Detecta si email ya tiene TKTs abiertos "
         "sin revelar si el email existe (solo indica estado abierto en esa empresa)."],
        ["ARCO-QW6: acuse de recibo automático", "ALTO", "routes/tkt_solicitud_derecho.py, services/email_service.py",
         "notificar_acuse_solicitud() al crear TKT con titular_email. Fallo registra warning, no revierte."],
        ["ARCO-QW7: chips de placeholders", "MEDIO", "frontend-next/components/tkt/TicketDrawer.tsx",
         "Chips de {{nombre_titular}}, {{empresa}}, {{fecha}}, {{numero_solicitud}}, {{dias_bloqueo}}, {{fecha_vencimiento}}."],
        ["ARCO-QW8: banner SLA en FlujoModal", "MEDIO", "frontend-next/components/arco/FlujoModal.tsx",
         "Semáforo verde/amarillo/rojo con días hábiles consumidos y restantes respecto al SLA legal (10 días hábiles Art. 12)."],
        ["Empresas-QW6: CompanyFichaPanel", "ALTO", "frontend-next/components/companies/CompanyFichaPanel.tsx",
         "Panel con 4 tabs lazy: Datos (edición), RATs, ARCO (tickets), Brechas."],
        ["CI/CD: pip-audit CVE scan", "ALTO", ".github/workflows/tests.yml",
         "CRITICAL bloquea deploy; HIGH genera advertencia. Env vars ALLOWED_ORIGINS y ENVIRONMENT hardcodeados para tests."],
        ["Nomenclatura: APDC → APDP", "MEDIO", "Global (código, tests, docs)",
         "Toda referencia a la agencia reguladora usa 'APDP' (Agencia de Protección de Datos Personales)."],
    ]
    add_table(doc, ["Cambio", "Severidad", "Archivos", "Descripción"], cambios_v112,
              col_widths_cm=[3.0, 1.5, 4.0, 9.5])

    dest = rename_to_v113(base)
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# MTX — Matriz de Trazabilidad
# ══════════════════════════════════════════════════════════════════════════════

def consolidar_mtx():
    print("\n[MTX] Matriz de Trazabilidad...")

    # Base: v1.7 — tiene tabla legal, hallazgos DTs, features→tests hasta Sprint 2
    ok = run_script(AUD_V17 / "build_MTX_matriz_v1_7.py")
    if not ok:
        return

    base = find_docx("Matriz_Trazabilidad")
    if not base:
        print("  WARN: no se encontró Matriz_Trazabilidad*.docx")
        return

    doc = Document(str(base))

    # v1.9: RF-163 a RF-169
    section_heading(doc, "Trazabilidad RF→HU→CU→TC — v1.9 (Julio 2026)", level=1)
    mtx_v19 = [
        ["RF-163", "HU-098", "CU-078", "TC-039–041, TC-043", "IDOR multi-tenant en 6 endpoints RAT (404 si empresa no coincide)", "CRÍTICO"],
        ["RF-164", "HU-099", "CU-079", "TC-044", "base_legal_valida strict contra enum taxativo (6 opciones)", "ALTO"],
        ["RF-165", "HU-100", "CU-080", "TC-045", "ConsentimientoAlert: listarConsentimientos() si datos_sensibles=True", "ALTO"],
        ["RF-166", "HU-101", "CU-081", "N/A", "Homologación orden campos RAT (5 pasos canónicos)", "ALTO"],
        ["RF-167", "HU-102", "CU-082", "TC-046", "PDF con títulos de sección y alertas rojas", "MEDIO"],
        ["RF-168", "HU-103", "CU-082", "N/A", "Encoding UTF-8 corregido en backend", "MEDIO"],
        ["RF-169", "N/A", "N/A", "N/A", "Código muerto eliminado (return duplicado, model_dump duplicado)", "BAJA"],
    ]
    add_table(doc, ["RF", "HU", "CU", "TC", "Descripción", "Severidad"], mtx_v19,
              col_widths_cm=[1.5, 1.5, 1.5, 2.0, 7.5, 2.0])

    # v1.12: RF-174 a RF-179, RNF-21, RNF-22
    section_heading(doc, "Trazabilidad RF→HU→CU→TC — v1.12 (Septiembre 2026)", level=1)
    mtx_v112 = [
        ["RF-174", "HU-104", "CU-083", "TC-056–059", "GET /publico/verificar-titular: detecta titular con TKTs abiertos (Art. 12)", "MEDIO"],
        ["RF-175", "HU-105", "CU-084", "TC-060–061", "POST TKT: acuse de recibo automático al titular con tracking token", "ALTO"],
        ["RF-176", "HU-106", "CU-085", "—", "TicketDrawer: chips de placeholders dinámicos en respuesta ARCO", "MEDIO"],
        ["RF-177", "HU-107", "CU-086", "—", "FlujoModal: semáforo SLA con días hábiles consumidos/restantes", "MEDIO"],
        ["RF-178", "HU-108", "CU-087", "TC-062", "CompanyFichaPanel: ficha empresa con 4 tabs lazy (Art. 16)", "ALTO"],
        ["RF-179", "—", "—", "—", "CI/CD pip-audit: cero vulnerabilidades CRITICAL en dependencias Python", "ALTO"],
        ["RNF-21", "—", "—", "TC-063", "Nomenclatura APDP correcta en código, tests y docs", "MEDIO"],
        ["RNF-22", "—", "—", "TC-059", "Rate limit 20/min/IP en GET /publico/verificar-titular", "MEDIO"],
    ]
    add_table(doc, ["RF", "HU", "CU", "TC", "Descripción", "Severidad"], mtx_v112,
              col_widths_cm=[1.5, 1.5, 1.5, 2.0, 7.5, 2.0])

    section_heading(doc, "Cobertura total v1.13", level=1)
    doc.add_paragraph(
        "Total RF documentados: 179 (RF-001 a RF-179) + 22 RNF (RNF-01 a RNF-22). "
        "Total HU: 108 (HU-001 a HU-108). Total CU: 87 (CU-001 a CU-087). "
        "Total TC: 63 (TC-001 a TC-063). RF con HU: 98%. RF con CU: 96%. RF con TC: 82%."
    )

    dest = rename_to_v113(base)
    if dest.exists():
        dest.unlink()
    doc.save(str(dest))
    if base != dest and base.exists():
        base.unlink()
    print(f"  ✓ {dest.name}")


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print(f"=== Generación consolidada {VERSION} ===")
    print(f"OUT_DIR: {OUT_DIR}")

    if not OUT_DIR.exists():
        print(f"ERROR: {OUT_DIR} no existe")
        sys.exit(1)

    consolidar_02_requisitos()
    consolidar_03_historias()
    consolidar_04_casos_uso()
    consolidar_08_api()
    consolidar_09_backlog()
    consolidar_10_plan_qa()
    consolidar_12_manual_tecnico()
    consolidar_mtx()

    print(f"\n✓ Consolidación {VERSION} completada.")
    print("Archivos en docs/documentacion_oficial/:")
    for f in sorted(OUT_DIR.glob("*.docx")):
        print(f"  {f.name}")
