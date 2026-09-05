"""
scripts/maintenance/generar_casos_uso_v1_10.py
Genera 04_Casos_de_Uso_Custodio_RAT_Manager_v1.10.docx agregando CU-15 a CU-25
(casos de uso de export, dashboard, eipd workflow).

Ejecutar:
  cd backend && python ../scripts/maintenance/generar_casos_uso_v1_10.py
"""

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor


CASOS_DE_USO = [
    # (id, titulo, actor, precondicion, flujo, postcondicion)
    ("CU-01", "Crear RAT basico",                "Admin Empresa / Editor",
     "Empresa existe, usuario autenticado con rol editor.",
     "1. Click en 'Crear RAT'\n2. Wizard 5 pasos\n3. Completar obligatorios Art. 16 (7 campos)\n4. Submit",
     "RAT creado en estado 'completo' o 'borrador' segun completitud."),

    ("CU-02", "Crear RAT con datos sensibles",   "Admin Empresa / Editor",
     "Empresa existe, usuario autenticado.",
     "1. Wizard paso 2 marca 'datos_sensibles=True'\n2. Especificar tipo_dato_sensible (Art. 2 g)\n3. Crear EIPD\n4. Completar EIPD\n5. Registrar consentimiento expreso",
     "RAT con sensibles + EIPD completada + consentimiento activo."),

    ("CU-03", "Crear RAT con transferencia internacional", "Admin Empresa / Editor",
     "Empresa existe.",
     "1. Wizard paso 4 marca 'transferencia_internacional=True'\n2. Especificar pais_destino\n3. Documentar garantias (SCC, BCR u otras)",
     "RAT con transferencia internacional + garantias documentadas."),

    ("CU-04", "Aprobar RAT",                     "Admin Empresa",
     "RAT en estado 'completo' con 100% completitud.",
     "1. GET /rats/{id}\n2. Verificar 100%\n3. POST /rats/{id}/aprobar",
     "RAT en estado 'aprobado' con aprobado_por + fecha_aprobacion."),

    ("CU-05", "Editar RAT",                      "Admin Empresa / Editor",
     "RAT existe, usuario tiene acceso.",
     "1. GET /rats/{id}\n2. PUT /rats/{id} con cambios\n3. Validadores condicionales",
     "RAT actualizado + audit log + alertas regeneradas."),

    ("CU-06", "Eliminar RAT",                    "Admin Empresa / Editor",
     "RAT existe, usuario tiene acceso.",
     "1. DELETE /rats/{id}\n2. Archivo movido a archive bucket",
     "RAT eliminado + audit log + archivo en archive."),

    ("CU-07", "Listar RATs con filtros",         "Cualquier rol con acceso",
     "Empresa existe.",
     "1. GET /rats/reportes?company_id=X&datos_sensibles=true&skip=0&limit=50\n2. Backend filtra + pagina",
     "Lista paginada con filtros aplicados."),

    ("CU-08", "Generar dashboard stats",         "Admin Empresa / Superadmin",
     "Empresa existe, usuario con acceso.",
     "1. GET /rats/dashboard/{company_id}\n2. Backend calcula KPIs",
     "JSON con KPIs: total, por_estado, sensibles, EIPD, transferencias, IL sin test, encargados sin contrato."),

    ("CU-09", "Exportar CSV",                    "Admin Empresa / Superadmin",
     "Empresa existe, usuario tiene acceso.",
     "1. GET /rats/export/csv?company_id=X\n2. Backend sanitiza CSV injection\n3. UTF-8 BOM para Excel",
     "Descarga archivo CSV con todos los RATs de la empresa."),

    ("CU-10", "Exportar PDF",                    "Admin Empresa / Superadmin",
     "Empresa existe, usuario tiene acceso.",
     "1. GET /rats/export/pdf?company_id=X\n2. Backend genera PDF con ReportLab\n3. Marca 'RAT BLOQUEADO' en rojo si aplica",
     "Descarga archivo PDF con todos los RATs."),

    ("CU-11", "Exportar RAT individual PDF",     "Admin Empresa / Superadmin",
     "RAT existe, usuario tiene acceso.",
     "1. GET /rats/{id}/export/pdf\n2. IDOR prevention via get_rat_for_user",
     "Descarga PDF del RAT individual."),

    ("CU-12", "Exportar formato CNI (APDC)",     "Admin Empresa / Superadmin",
     "Empresa existe.",
     "1. GET /rats/export/cni?company_id=X\n2. Backend serializa en formato APDC Ley 21.719",
     "Descarga archivo .txt en formato APDC para presentar ante regulador."),

    ("CU-13", "Registrar consentimiento expreso", "Admin Empresa / Editor",
     "RAT con datos_sensibles=True.",
     "1. POST /rats/{id}/consentimientos\n2. Backend cifra PII con Fernet\n3. SHA-256 sobre texto del consentimiento",
     "Consentimiento activo vinculado al RAT."),

    ("CU-14", "Revocar consentimiento",          "Admin Empresa / Editor",
     "Consentimiento activo existe.",
     "1. POST /consentimientos/{id}/revocar\n2. activo=False + fecha_revocacion",
     "Consentimiento revocado, no más tratamientos."),

    # CU-15 a CU-25 — nuevos en v1.10
    ("CU-15", "Paginacion de reportes >100 RATs", "Cualquier rol con acceso",
     "Empresa tiene 100+ RATs registrados.",
     "1. GET /rats/reportes?company_id=X&skip=0&limit=50\n2. Backend pagina resultados\n3. Cliente navega con skip/limit",
     "Pagina de 50 RATs. Cliente itera skip hasta total/skipped."),

    ("CU-16", "Auditoria del RAT",                "Cualquier rol con acceso",
     "RAT existe.",
     "1. GET /rats/{id}/auditoria\n2. Backend valida pertenencia via get_rat_for_user\n3. Retorna historial",
     "Lista de audit logs del RAT con hash chain."),

    ("CU-17", "Auditoria global empresa",         "Admin Empresa / Superadmin",
     "Empresa existe.",
     "1. GET /rats/auditoria/{company_id}\n2. Backend filtra audit logs por rat_ids de la empresa",
     "Lista paginada de todos los eventos de la empresa."),

    ("CU-18", "Verificar integridad cadena",     "Solo SUPERADMIN",
     "Usuario es SUPERADMIN.",
     "1. GET /rats/auditoria/verify-chain\n2. Backend verifica SHA-256 chain\n3. Retorna valido/error_id",
     "Resultado de verificacion de cadena de hashes."),

    ("CU-19", "Descargar archivo base legal",    "Editor / Admin Empresa",
     "RAT tiene archivo adjunto.",
     "1. GET /rats/{id}/archivo\n2. Backend intenta OCI → fallback BYTEA descifrado\n3. Retorna bytes o presigned URL",
     "Archivo PDF descifrado descargado."),

    ("CU-20", "Sugerencias automaticas por rubro", "Cualquier rol",
     "Empresa tiene rubro configurado.",
     "1. GET /rats/sugerencias/tipos → lista tipos\n2. POST /rats/sugerencias con tipo_proceso\n3. Backend retorna sugerencia precompletada",
     "Wizard pre-rellenado con sugerencia del backend."),

    ("CU-21", "Registrar revision periodica",    "Editor / Admin Empresa",
     "RAT existe.",
     "1. POST /rats/{id}/revision\n2. Backend actualiza updated_at + audit log",
     "RAT marcado como revisado."),

    ("CU-22", "Buscar RATs por texto",           "Cualquier rol con acceso",
     "Empresa existe.",
     "1. GET /rats/reportes?company_id=X&search=marketing\n2. Backend filtra por nombre_proceso (ilike)",
     "Lista de RATs que coinciden con el termino."),

    ("CU-23", "Filtros combinados",              "Cualquier rol con acceso",
     "Empresa existe.",
     "1. GET /rats/reportes?transferencia_internacional=true&decisiones_automatizadas=true&datos_sensibles=true\n2. Backend aplica AND",
     "RATs que cumplen TODOS los filtros."),

    ("CU-24", "Duplicar RAT",                    "Editor / Admin Empresa",
     "RAT existe.",
     "1. Click en icono duplicar\n2. Frontend copia campos\n3. POST /rats/ con datos copiados\n4. Nuevo RAT independiente",
     "Nuevo RAT con datos similares, ID diferente."),

    ("CU-25", "Bloquear RAT temporalmente",      "Sistema (Art. 8 ter)",
     "RAT existe, incidente detectado.",
     "1. PUT /rats/{id} con bloqueado=true\n2. Export PDF/CSV marca en rojo 'RAT BLOQUEADO'",
     "RAT bloqueado visualmente en exports."),
]


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for level, size, color in [
        ('Heading 1', 16, RGBColor(0x1F, 0x49, 0x7D)),
        ('Heading 2', 14, RGBColor(0x2E, 0x74, 0xB5)),
        ('Heading 3', 12, RGBColor(0x40, 0x40, 0x40)),
    ]:
        s = doc.styles[level]
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color


def add_cu_table(doc, casos):
    """Tabla resumen de casos de uso."""
    headers = ['ID', 'Titulo', 'Actor']
    table = doc.add_table(rows=1 + len(casos), cols=3)
    table.style = 'Light Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for r, (cid, titulo, actor, *_rest) in enumerate(casos, start=1):
        row = table.rows[r].cells
        row[0].text = cid
        row[1].text = titulo
        row[2].text = actor


def add_cu_detail(doc, cid, titulo, actor, precondicion, flujo, postcondicion):
    doc.add_heading(f"{cid}: {titulo}", level=2)
    p = doc.add_paragraph()
    p.add_run("Actor: ").bold = True
    p.add_run(actor)
    p = doc.add_paragraph()
    p.add_run("Precondicion: ").bold = True
    p.add_run(precondicion)
    p = doc.add_paragraph()
    p.add_run("Flujo: ").bold = True
    p.add_run(flujo)
    p = doc.add_paragraph()
    p.add_run("Postcondicion: ").bold = True
    p.add_run(postcondicion)


def main():
    script_dir = Path(__file__).resolve().parent.parent.parent
    docs_dir = script_dir / "docs" / "documentacion_oficial"
    output = docs_dir / "04_Casos_de_Uso_Custodio_RAT_Manager_v1.10.docx"

    doc = Document()
    setup_styles(doc)

    doc.add_heading('CUSTODIO RAT — Casos de Uso v1.10', 0)
    doc.add_paragraph('Sistema RAT Manager — Ley 21.719 de Proteccion de Datos Personales de Chile')

    doc.add_heading('Metadata', 1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid'
    headers = ['Version', 'Fecha', 'Autor', 'Cambios']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    row = table.rows[1].cells
    row[0].text = '1.10'
    row[1].text = '2026-07-07'
    row[2].text = 'Auditoria RAT detallada — Custodio'
    row[3].text = f'{len(CASOS_DE_USO)} casos de uso (antes 14 en v1.9). Agregados CU-15 a CU-25 (export, dashboard, paginacion, duplicacion, bloqueo).'

    doc.add_heading('Resumen de casos de uso', 1)
    doc.add_paragraph(f'Total: {len(CASOS_DE_USO)} casos de uso.')
    add_cu_table(doc, CASOS_DE_USO)

    doc.add_page_break()

    doc.add_heading('Detalle de casos de uso', 1)
    for cu in CASOS_DE_USO:
        add_cu_detail(doc, *cu)

    docs_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"OK: Generado {output}")
    print(f"   Casos de uso: {len(CASOS_DE_USO)}")


if __name__ == "__main__":
    main()